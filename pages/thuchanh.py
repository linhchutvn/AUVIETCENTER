import streamlit as st
from google import genai
from google.genai import types
import json
import re
import time
import random
import textwrap
import html
import os
import requests
from PIL import Image
from io import BytesIO

# Thư viện Word
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Thư viện PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.fonts import addMapping

# ==========================================
# 1. CẤU HÌNH TRANG (PHẢI ĐẶT ĐẦU TIÊN)
# ==========================================
st.set_page_config(page_title="IELTS Writing Master", page_icon="🎓", layout="wide")

# ==========================================
# 2. CSS TỔNG HỢP (ẨN HEADER/FOOTER + STYLE APP)
# ==========================================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Merriweather:wght@300;400;700&display=swap');
    
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; }

    /* --- PHẦN ẨN GIAO DIỆN MẶC ĐỊNH --- */
    
    /* 1. Ẩn thanh Header trên cùng (Chứa nút 3 chấm và Running man) */
    .stAppHeader {
        display: none;
    }
    
    /* 2. Ẩn Footer 'Made with Streamlit' */
    footer {
        visibility: hidden;
    }
    
    /* 3. Ẩn nút Deploy (Con thuyền màu đỏ) */
    .stDeployButton {
        display: none;
    }
    
    /* 4. Ẩn Menu Hamburger (nếu CSS trên chưa ẩn hết) */
    #MainMenu {
        visibility: hidden;
    }

    /* --- PHẦN STYLE GIAO DIỆN APP --- */
    
    /* Header Style */
    .main-header {
        font-family: 'Merriweather', serif;
        color: #0F172A;
        font-weight: 700;
        font-size: 2.2rem;
        margin-bottom: 0rem;
        margin-top: -2rem; /* Đẩy tiêu đề lên cao hơn vì đã ẩn Header */
    }
    .sub-header {
        font-family: 'Inter', sans-serif;
        color: #64748B;
        font-size: 1.1rem;
        margin-bottom: 1rem;
        border-bottom: 1px solid #E2E8F0;
        padding-bottom: 0.5rem;
    }

    /* Step Headers */
    .step-header {
        font-family: 'Inter', sans-serif;
        font-weight: 700;
        font-size: 1.2rem;
        color: #1E293B;
        margin-top: 1.5rem;
        margin-bottom: 0.5rem;
    }
    .step-desc {
        font-size: 0.9rem;
        color: #64748B;
        margin-bottom: 0.8rem;
    }
    /* --- ẨN CÁC ICON GHIM (LINK CHAIN) BÊN CẠNH TIÊU ĐỀ --- */
    [data-testid="stMarkdownContainer"] h1 a,
    [data-testid="stMarkdownContainer"] h2 a,
    [data-testid="stMarkdownContainer"] h3 a,
    [data-testid="stMarkdownContainer"] h4 a,
    [data-testid="stMarkdownContainer"] h5 a,
    [data-testid="stMarkdownContainer"] h6 a {
        display: none !important;
        pointer-events: none;
    }

    /* Guide Box */
    .guide-box {
        background-color: #f8f9fa;
        border-left: 5px solid #ff4b4b;
        padding: 15px;
        border-radius: 5px;
        margin-bottom: 10px;
        color: #31333F;
    }

    /* Error Cards */
    .error-card {
        background-color: white;
        border: 1px solid #E5E7EB;
        border-radius: 12px;
        padding: 20px;
        margin-bottom: 16px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
        transition: all 0.2s;
    }
    .error-card:hover {
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
        border-color: #D1D5DB;
    }
    
    .annotated-text {
        font-family: 'Merriweather', serif;
        line-height: 1.8;
        color: #374151;
        background-color: white;
        padding: 24px;
        border-radius: 12px;
        border: 1px solid #E5E7EB;
        box-shadow: 0 1px 2px rgba(0,0,0,0.05);
    }
    
    del { color: #9CA3AF; text-decoration: line-through; margin-right: 4px; text-decoration-thickness: 2px; }
    ins.grammar { background-color: #4ADE80; color: #022C22; text-decoration: none; padding: 2px 6px; border-radius: 4px; font-weight: 700; border: 1px solid #22C55E; }
    ins.vocab { background-color: #FDE047; color: #000; text-decoration: none; padding: 2px 6px; border-radius: 4px; font-weight: 700; border: 1px solid #FCD34D; }
    
    /* Button Customization */
    div.stButton > button {
        background-color: #FF4B4B;
        color: white;
        font-weight: bold;
        border-radius: 8px;
        padding: 0.5rem 1.5rem;
        border: none;
    }
    div.stButton > button:hover {
        background-color: #D93434;
    }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 2. LOGIC AI (FAILOVER)
# ==========================================
try:
    ALL_KEYS = st.secrets["GEMINI_API_KEYS"]
except Exception:
    st.error("⚠️ Chưa cấu hình secrets.toml chứa GEMINI_API_KEYS!")
    st.stop()

import streamlit as st
from google import genai
from google.genai import types
import json
import re
import time
import random
import textwrap
import html
import os
import requests
from PIL import Image
from io import BytesIO

# Thư viện Word
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Thư viện PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.fonts import addMapping

# ==========================================
# 1. CẤU HÌNH TRANG (PHẢI ĐẶT ĐẦU TIÊN)
# ==========================================
st.set_page_config(page_title="IELTS Writing Master", page_icon="🎓", layout="wide")

# ==========================================
# 2. CSS TỔNG HỢP (ẨN HEADER/FOOTER + STYLE APP)
# ==========================================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Merriweather:wght@300;400;700&display=swap');
    
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; }

    /* --- PHẦN ẨN GIAO DIỆN MẶC ĐỊNH --- */
    
    /* 1. Ẩn thanh Header trên cùng (Chứa nút 3 chấm và Running man) */
    .stAppHeader {
        display: none;
    }
    
    /* 2. Ẩn Footer 'Made with Streamlit' */
    footer {
        visibility: hidden;
    }
    
    /* 3. Ẩn nút Deploy (Con thuyền màu đỏ) */
    .stDeployButton {
        display: none;
    }
    
    /* 4. Ẩn Menu Hamburger (nếu CSS trên chưa ẩn hết) */
    #MainMenu {
        visibility: hidden;
    }

    /* --- PHẦN STYLE GIAO DIỆN APP --- */
    
    /* Header Style */
    .main-header {
        font-family: 'Merriweather', serif;
        color: #0F172A;
        font-weight: 700;
        font-size: 2.2rem;
        margin-bottom: 0rem;
        margin-top: -2rem; /* Đẩy tiêu đề lên cao hơn vì đã ẩn Header */
    }
    .sub-header {
        font-family: 'Inter', sans-serif;
        color: #64748B;
        font-size: 1.1rem;
        margin-bottom: 1rem;
        border-bottom: 1px solid #E2E8F0;
        padding-bottom: 0.5rem;
    }

    /* Step Headers */
    .step-header {
        font-family: 'Inter', sans-serif;
        font-weight: 700;
        font-size: 1.2rem;
        color: #1E293B;
        margin-top: 1.5rem;
        margin-bottom: 0.5rem;
    }
    .step-desc {
        font-size: 0.9rem;
        color: #64748B;
        margin-bottom: 0.8rem;
    }
    /* --- ẨN CÁC ICON GHIM (LINK CHAIN) BÊN CẠNH TIÊU ĐỀ --- */
    [data-testid="stMarkdownContainer"] h1 a,
    [data-testid="stMarkdownContainer"] h2 a,
    [data-testid="stMarkdownContainer"] h3 a,
    [data-testid="stMarkdownContainer"] h4 a,
    [data-testid="stMarkdownContainer"] h5 a,
    [data-testid="stMarkdownContainer"] h6 a {
        display: none !important;
        pointer-events: none;
    }

    /* Guide Box */
    .guide-box {
        background-color: #f8f9fa;
        border-left: 5px solid #ff4b4b;
        padding: 15px;
        border-radius: 5px;
        margin-bottom: 10px;
        color: #31333F;
    }

    /* Error Cards */
    .error-card {
        background-color: white;
        border: 1px solid #E5E7EB;
        border-radius: 12px;
        padding: 20px;
        margin-bottom: 16px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
        transition: all 0.2s;
    }
    .error-card:hover {
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
        border-color: #D1D5DB;
    }
    
    .annotated-text {
        font-family: 'Merriweather', serif;
        line-height: 1.8;
        color: #374151;
        background-color: white;
        padding: 24px;
        border-radius: 12px;
        border: 1px solid #E5E7EB;
        box-shadow: 0 1px 2px rgba(0,0,0,0.05);
    }
    
    del { color: #9CA3AF; text-decoration: line-through; margin-right: 4px; text-decoration-thickness: 2px; }
    ins.grammar { background-color: #4ADE80; color: #022C22; text-decoration: none; padding: 2px 6px; border-radius: 4px; font-weight: 700; border: 1px solid #22C55E; }
    ins.vocab { background-color: #FDE047; color: #000; text-decoration: none; padding: 2px 6px; border-radius: 4px; font-weight: 700; border: 1px solid #FCD34D; }
    
    /* Button Customization */
    div.stButton > button {
        background-color: #FF4B4B;
        color: white;
        font-weight: bold;
        border-radius: 8px;
        padding: 0.5rem 1.5rem;
        border: none;
    }
    div.stButton > button:hover {
        background-color: #D93434;
    }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 2. LOGIC AI (FAILOVER)
# ==========================================
ALL_KEYS = st.secrets["GEMINI_API_KEYS"]

def generate_content_with_failover(prompt, image=None, json_mode=False):
    import time  # Đảm bảo đã import time
    
    keys_to_try = list(ALL_KEYS)
    random.shuffle(keys_to_try) 
    
    model_priority = [
        #"gemini-3-flash-preview",        
        "gemini-2.5-flash",
        "gemini-2.5-flash-lite",
        "gemini-2.0-flash",
        "gemini-1.5-pro", 
        "gemini-1.5-flash"
    ]
    
    last_error = ""
    # 💡 BỔ SUNG: Khởi tạo vùng thông báo để không bị lỗi NameError
    status_msg = st.empty() 

    for index, current_key in enumerate(keys_to_try):
        try:
            # --- BƯỚC 1: Khởi tạo kết nối & Né chặn IP ---
            if index > 0:
                status_msg.warning(f"⏳ Luồng #{index} bận. Đang tối ưu kết nối, vui lòng đợi 3 giây...")
                time.sleep(3) 
            
            client = genai.Client(api_key=current_key)
            
            # --- BƯỚC 2: Lấy danh sách model ---
            raw_models = list(client.models.list())
            available_models = [m.name.replace("models/", "") for m in raw_models]
            
            # --- BƯỚC 3: Tìm model tốt nhất ---
            sel_model = None
            for target in model_priority:
                if target in available_models:
                    sel_model = target
                    break
            
            if not sel_model:
                sel_model = "gemini-1.5-flash" 

            # --- BƯỚC 4: Hiển thị thông tin Debug ---
            masked_key = f"****{current_key[-4:]}"
            st.toast(f"⚡ Connected: {sel_model}", icon="🤖")
            
            with st.expander(f"🔌 Connection Details (Key #{index + 1})", expanded=False):
                st.write(f"**Active Model:** `{sel_model}`")
                st.write(f"**Active API Key:** `{masked_key}`")
            
            # --- BƯỚC 5: Chuẩn bị nội dung ---
            content_parts = [image, prompt] if image else [prompt]
                
            # --- BƯỚC 6: Cấu hình ---
            config_args = {
                "temperature": 0.3,
                "top_p": 0.95,
                "top_k": 64,
                "max_output_tokens": 32000,
            }
            
            if json_mode and "thinking" not in sel_model.lower():
                config_args["response_mime_type"] = "application/json"

            if "thinking" in sel_model.lower():
                config_args["thinking_config"] = {"include_thoughts": True, "thinking_budget": 32000}

            # --- BƯỚC 7: Thực hiện gọi API ---
            # Xóa thông báo chờ trước khi gọi AI
            status_msg.info(f"🚀 Processing data via Stream #{index + 1}...")
            
            response = client.models.generate_content(
                model=sel_model,
                contents=content_parts,
                config=types.GenerateContentConfig(**config_args)
            )
            
            status_msg.empty() # Thành công thì xóa thông báo
            return response, sel_model 
            
        except Exception as e:
            last_error = str(e)
            if "429" in last_error or "quota" in last_error.lower():
                continue 
            else:
                st.warning(f"⚠️ Luồng #{index+1} gặp sự cố kỹ thuật. Đang chuyển luồng...")
                continue
                
    status_msg.empty()
    st.error(f"❌ Tất cả {len(keys_to_try)} luồng kết nối đều thất bại. Vui lòng thử lại sau 1 phút.")
    return None, None

# ==========================================
# 3. PROMPT KHỦNG (NGUYÊN BẢN TỪ APP CHẤM ĐIỂM)
# ==========================================
GRADING_PROMPT_TEMPLATE = """
Bạn hãy đóng vai trò là một Giám khảo IELTS với 30 năm kinh nghiệm làm việc tại Hội đồng Anh (British Council). Nhiệm vụ của bạn là đánh giá bài viết dựa trên **bộ tiêu chí chuẩn xác của IELTS Writing Task 1 (Band Descriptors)**. 
**Phân loại bài thi (Context Awareness):** Bắt buộc phải nhận diện đây là IELTS Academic: Biểu đồ/Đồ thị/Quy trình/Map. Đề bài nói về nội dung gì.
**Yêu cầu khắt khe:** Bạn phải sử dụng **tiêu chuẩn của Band 9.0 làm thước đo tham chiếu cao nhất** để soi xét bài làm. Hãy thực hiện một bản "Gap Analysis" chi tiết: chỉ ra mọi thiếu sót một cách nghiêm ngặt và chính xác tuyệt đối, từ những lỗi sai căn bản cho đến những điểm chưa đạt được độ tinh tế của một bài viết điểm tuyệt đối.
**YÊU CẦU ĐẶC BIỆT (CHẾ ĐỘ KIỂM TRA KỸ):** Bạn không cần phải trả lời nhanh. Hãy dành thời gian "suy nghĩ" để phân tích thật sâu và chi tiết (Step-by-step Analysis).

### 1. TƯ DUY & GIAO THỨC LÀM VIỆC (CORE PROTOCOL)
* **>> GIAO THỨC PHÂN TÍCH CHẬM (SLOW REASONING PROTOCOL):**
    * Bạn không được phép tóm tắt nhận xét. Với mỗi tiêu chí, bạn phải viết ít nhất 200-300 từ.
    * Bạn phải thực hiện phân tích theo phương pháp "Socratic": Đặt câu hỏi về từng câu văn của thí sinh, tìm ra điểm chưa hoàn hảo và giải thích cặn kẽ tại sao nó chưa đạt Band 7.0 hoặc Band 9.0 từ dữ liệu bài viết này.
    * Cấm dùng các cụm từ chung chung như "Good grammar" hay "Appropriate vocabulary". Bạn phải trích dẫn ít nhất 3-5 ví dụ thực tế từ bài làm cho mỗi tiêu chí để chứng minh cho nhận định của mình.
*   **Persona:** Giám khảo lão làng, khó tính nhưng công tâm. Tông giọng phản hồi trực diện, không khen ngợi sáo rỗng. Nếu bài tệ, phải nói rõ là tệ.
*   **>> NGUYÊN TẮC "HOLISTIC SCORING" (Chấm điểm tổng hòa):** 
    *   Tuyệt đối phân biệt giữa **Lỗi hệ thống (Systematic error)** và **Lỗi trượt chân (Slip)**.
    *   *Lỗi trượt chân (Slip):* Là lỗi nhỏ, ngẫu nhiên (như viết thiếu 1 chữ cái, thừa 1 từ so sánh). Nếu bài viết thể hiện trình độ từ vựng/ngữ pháp xuất sắc, những lỗi này **KHÔNG ĐƯỢC** dùng làm lý do để hạ điểm từ 8 xuống 7 hoặc từ 9 xuống 8.
*   **Chế độ "Deep Scan":** Không trả lời nhanh. Hãy dành thời gian phân tích từng câu, từng từ theo quy trình "Step-by-step Analysis".
*   **Quy tắc "Truy quét kiệt quệ" (Exhaustive Listing):**
    *   Tuyệt đối KHÔNG gộp lỗi. Nếu thí sinh sai 10 lỗi mạo từ, liệt kê đủ 10 mục.
    *   Danh sách lỗi trong JSON là bằng chứng pháp lý. Mọi lỗi nhỏ nhất (dấu phẩy, viết hoa, mạo từ) đều phải được ghi nhận. Nếu JSON ít lỗi mà điểm GRA thấp, đó là một sự mâu thuẫn nghiêm trọng.
    *   **>> BỔ SUNG QUY TẮC TAXONOMY:** Khi phân loại lỗi trong JSON, chỉ được sử dụng các thuật ngữ chuẩn mực (ví dụ: Subject-Verb Agreement, Collocation, Article, Comma Splice). TUYỆT ĐỐI KHÔNG sáng tạo ra tên lỗi lạ (như "Bad word", "Wrong grammar").
*   **Nhận diện ngữ cảnh (Context Awareness):** Tự xác định là Academic (Biểu đồ/Process/Map) hay General Training (Thư) để áp dụng Band Descriptors tương ứng.
* **>> GIAO THỨC QUÉT 2 LỚP (TWO-PASS SCANNING):**
    * Lớp 1: Tìm các lỗi nặng (Cấu trúc, từ vựng sai ngữ cảnh, logic dữ liệu).
    * Lớp 2: Quét lại toàn bộ bài để tìm các lỗi nhỏ (Mạo từ, số ít/nhiều, dấu câu, viết hoa). 
    * Chỉ sau khi hoàn thành 2 lớp quét này mới được lập danh sách lỗi cuối cùng.
*   **>> NGUYÊN TẮC "APPROXIMATION TOLERANCE":** 
    *   Đối với các số liệu rất nhỏ (< 2-3%), chấp nhận các từ ngữ ước lượng mạnh như *"virtually no"*, *"almost zero"*, *"negligible"*. Đừng coi đây là lỗi sai dữ liệu (Logic Error) trừ khi số liệu thực tế > 5%.    

### 2. TIÊU CHÍ CHẤM ĐIỂM CHI TIẾT (4 CRITERIA)
#### A. Task Achievement (TA)
*   **Tư duy dữ liệu & Nhóm thông tin (Logical Grouping):**
    *   **Band 8.0+:** Thí sinh PHẢI biết nhóm các đối tượng tương đồng vào cùng đoạn văn một cách thông minh (Skilfully selected). Nếu chỉ liệt kê máy móc -> Tối đa Band 6-7.
    *   **>> BỔ SUNG QUY TẮC CHẶN BAND 6 (Comparison Rule):** Nếu bài viết chỉ mô tả đơn lẻ (description) số liệu của từng đối tượng mà KHÔNG CÓ sự so sánh (comparison) tương quan giữa các đối tượng -> **TỐI ĐA BAND 6.0** (Dù mô tả đúng 100%).
    *   **>> BỔ SUNG QUY TẮC "TOTAL/OTHER" (Safety Net):** Các hạng mục như 'Total', 'Miscellaneous', 'Other' KHÔNG ĐƯỢC tính là Key Features bắt buộc. Nếu thí sinh bỏ qua các số liệu này, HOÀN TOÀN KHÔNG ĐƯỢC TRỪ ĐIỂM. (Cảnh báo: Nếu trừ điểm lỗi này là sai quy chế).
*   **Độ dài & Sự súc tích (Word Count vs Conciseness):**
    *   **Không phạt oan:** Nếu bài > 200 từ nhưng thông tin đắt giá, số liệu chính xác 100% -> KHÔNG hạ điểm TA.
    *   `>> ƯU TIÊN "DATA SYNTHESIZING": Đánh giá cao nếu thí sinh biết biến số liệu % thành phân số (fractions) hoặc các cụm từ ước lượng (rounding) thay vì chỉ liệt kê số liệu thô từ bảng.`
    *   **Chỉ trừ điểm khi:** Bài viết dài dòng do lặp ý (Repetitive) hoặc lan man (Irrelevant). Nếu > 200 từ mà nội dung tốt, chỉ đưa vào phần "Lời khuyên" là nên cô đọng hơn.
*   **>> QUY TẮC XỬ LÝ ĐỘ DÀI (WORD COUNT THRESHOLDS):**
    *   **Nguyên tắc cốt lõi:** Không trừ điểm chỉ vì con số, hãy trừ điểm vì **HỆ QUẢ** của việc thiếu từ (thiếu chi tiết, thiếu so sánh).
    *   **Zone A (140 - 149 words):** 
        *   Chế độ: "Khoan hồng" (Leniency).
        *   Nếu bài viết vẫn đủ Overview, số liệu và so sánh -> **KHÔNG TRỪ ĐIỂM**. Vẫn có thể đạt Band 7-8.
        *   Chỉ trừ điểm nếu thấy nội dung bị cắt gọt quá đà.
    *   **Zone B (100 - 139 words):** 
        *   Chế độ: "Cảnh báo Đỏ" (Red Alert).
        *   Hệ quả: Thường dẫn đến lỗi *"Limited detail"* (Chi tiết hạn chế) hoặc *"Key features not fully covered"*.
        *   **Hành động:** Kiểm tra gắt gao. Nếu thiếu thông tin -> **Block ngay ở Band 5.0 - 5.5 TA**. Khó có thể lên Band 6.
    *   **Zone C (21 - 99 words):**
        *   Chế độ: "Trừng phạt" (Penalty).
        *   Hệ quả: Vi phạm tiêu chí Band 3 (*"Significantly underlength"*).
        *   **Hành động:** **TỐI ĐA BAND 3.0 - 4.0 TA**. Không cần xét đến chất lượng câu chữ.
    *   **Zone D (0 - 20 words):** 
        *   **Hành động:** **BAND 1.0** (Theo đúng Band Descriptors).
*   **Các bẫy "Chết người" (Negative Features - TA):**
    *   **Object vs Figure:** Phạt nặng lỗi sai chủ ngữ (VD: "The figure of apple rose" -> Sai; "The consumption of apple rose" -> Đúng).
    *   **Nhầm đơn vị:** Đề là % mà viết là Number -> Chặn đứng ở Band 5.0 TA.
    *   **No Data/Support:** Academic mà mô tả không có số liệu đi kèm -> Band 5.0.
    *   **Band 5 (Nguy hiểm):** Nếu mô tả xu hướng mà **không có số liệu (data)** đi kèm -> BẮT BUỘC hạ xuống Band 5 (Theo dòng in đậm: "There may be no data to support the description").
    *   **Overview:** Process phải đủ "Đầu-Giữa-Cuối"; Map phải có "Sự thay đổi tổng quan". Sai/Thiếu Overview -> Tối đa Band 5-6.
    *   **Band 7:** Phải xác định được xu hướng chính/sự khác biệt rõ ràng (Clear overview).
    *   **Band 6:** Có nỗ lực viết Overview nhưng thông tin chọn lọc sai hoặc không rõ ràng.
    *   **Band 5:** Không có Overview hoặc Overview sai lệch hoàn toàn.
    *   **Ý kiến cá nhân:** Tuyệt đối cấm. Có ý kiến cá nhân -> Trừ điểm nặng.
    *   **>> QUY TẮC "MISSING INTRODUCTION" (Lỗi Định dạng):**
        *   Kiểm tra câu đầu tiên của bài viết. Nếu thí sinh nhảy bổ vào mô tả xu hướng/số liệu (Overview/Body) mà KHÔNG CÓ câu giới thiệu chủ đề (Paraphrase đề bài) -> **TỐI ĐA BAND 5.0 TA** (Lỗi "Inappropriate format").
        *   **Lý do:** Người đọc không biết biểu đồ nói về cái gì = Mất ngữ cảnh giao tiếp (Failure in Communication).
    *   **>> QUY TẮC "COPIED RUBRIC" (Sao chép đề):**
        *   So sánh câu mở đầu với đề bài. Nếu giống > 80% (chép nguyên văn các chuỗi từ dài) -> Những từ này KHÔNG được tính vào độ dài bài viết và vốn từ vựng.
        *   Nếu cả bài chỉ dựa vào đề bài chép lại -> **BAND 1 (Wholly unrelated/Copied).**
*   **>> BỔ SUNG QUY TẮC FORMAT & TONE:**
        *   **Lỗi định dạng (Format):** Nếu bài viết dùng gạch đầu dòng (bullet points) hoặc đánh số (1, 2, 3) thay vì viết đoạn văn -> **TỐI ĐA BAND 5.0 TA**.
        *   **Lỗi giọng điệu (Tone - GT):** Nếu đề yêu cầu "Formal letter" mà dùng ngôn ngữ suồng sã (slang, contractions like "gonna") -> Trừ điểm nặng xuống **Band 5.0-6.0**.
*   **Math Logic Check:** Soi kỹ các từ chỉ mức độ (slight, significant). Ví dụ: Từ 10% lên 15% là tăng gấp rưỡi -> Cấm dùng "slight".
*   **Endpoint Trap:** Cấm dùng "peak" cho năm cuối cùng của biểu đồ (vì không biết tương lai). Gợi ý: "ending at a high".
*   **>> CHIẾN THUẬT OVERVIEW BAND 8.0-9.0 (BẮT BUỘC ĐỐI CHIẾU):**
    1.  **Nguyên tắc "No Data":** Overview đạt Band cao TUYỆT ĐỐI không được chứa số liệu chi tiết. 
    2.  **Cấu trúc "Double Content":** Phải bao quát được cả (1) Xu hướng chính (Trends) VÀ (2) Sự so sánh nổi bật nhất (Major Comparisons/High-lows).
    3.  **Kỹ thuật Synthesis:** Đánh giá xem học sinh có biết gộp các đối tượng tương đồng để khái quát hóa không, hay chỉ đang liệt kê.
    4.  **Vị trí:** Khuyên học sinh đặt ngay sau Introduction để tạo luồng logic.
#### B. Coherence & Cohesion (CC)
*   **Liên kết "Vô hình" (Invisible Cohesion - Band 9):** Ưu tiên các cấu trúc "respectively", "in that order", mệnh đề quan hệ rút gọn.
*   **Mechanical Linkers (Lỗi máy móc):** Nếu câu nào cũng bắt đầu bằng "Firstly, Secondly, In addition, Furthermore" -> Tối đa Band 6.0.
*   **Paragraphing:** Bài viết phải chia đoạn logic. Chỉ có 1 đoạn văn -> CC tối đa 5.0.
*   **>> BỔ SUNG QUY TẮC "AMBIGUOUS REFERENCING" (The 'It' Trap):**
        *   Kiểm tra kỹ các đại từ thay thế (It, This, That, These, Those). Nếu dùng các từ này mà KHÔNG RÕ thay thế cho danh từ nào trước đó (gây khó hiểu) -> **TỐI ĐA BAND 6.0 CC**.
*   **>> QUY TẮC "INVISIBLE GLUE" (Keo dán vô hình):**
        *   Soi kỹ các từ dẫn đầu đoạn (Signposting words). Nếu thí sinh dùng lặp lại các từ như "Regarding...", "As for...", "Turning to..." quá 2 lần -> Đánh dấu là "Mechanical" (Máy móc).
        *   Khuyến khích cách chuyển đoạn bằng chủ ngữ ẩn hoặc Reference (Ví dụ: Thay vì "Regarding A, it increased...", hãy viết "A, conversely, witnessed a rise...").
*   **>> NGUYÊN TẮC LINH HOẠT CC:** Nếu bài viết có logic tốt và dễ hiểu, việc sử dụng từ nối hơi máy móc (như "Regarding") KHÔNG NÊN kéo điểm xuống 7.0 ngay lập tức. Hãy cân nhắc Band 8.0 nếu dòng chảy thông tin (flow) vẫn mượt mà. Chỉ hạ xuống 7.0 nếu việc dùng từ nối gây khó chịu hoặc làm gián đoạn việc đọc.
*   **>> YÊU CẦU OUTPUT CHO PHẦN NÀY:**
    *   **Trích dẫn chứng:** Phải trích dẫn câu văn cụ thể của thí sinh để phân tích.
    *   **Gợi ý "Vừa sức":** 
        *   Bài dưới Band 7 -> Gợi ý sửa cho ĐÚNG.
        *   Bài Band 7+ -> Gợi ý sửa cho HAY (Band 9).
#### C. Lexical Resource (LR)
*   **Naturalness over Academic:** Ưu tiên từ vựng tự nhiên (use, help, start) hơn là từ đao to búa lớn sai ngữ cảnh (utilise, facilitate, commence).
*   **Blacklist:** Cảnh báo các từ sáo rỗng/học thuộc lòng bị lạm dụng.
*   **Precision:** Soi kỹ Collocation (VD: "increased significantly" > "increased strongly").
*   **>> BỔ SUNG QUY TẮC "REPETITION" (Lặp từ):**
        *   Nếu một từ vựng quan trọng (ví dụ: "increase", "fluctuate") bị lặp lại > 3 lần mà không có nỗ lực thay thế (paraphrase) -> **TỐI ĐA BAND 5.0 LR** (Lỗi "Limited flexibility").
    *   **>> QUY TẮC CHÍNH TẢ (Spelling Threshold):**
        *   Sai 1-2 lỗi nhỏ -> Vẫn có thể Band 8.
        *   Sai vài lỗi (A few) nhưng vẫn hiểu được -> Band 7.
        *   Sai nhiều lỗi (Noticeable) nhưng vẫn hiểu được -> Band 6.
        *   Sai gây khó hiểu (Impede meaning) -> Band 5.
*   **>> NGUYÊN TẮC "NO DOUBLE PENALIZATION" (Không phạt kép):**
        *   Nếu lỗi thuộc về Redundancy (thừa từ: *most highest*) hoặc Spelling (*fluctation*), hãy tính nó vào điểm Lexical Resource (LR).
        *   KHÔNG trừ điểm Grammatical Range (GRA) cho những lỗi đã tính ở LR, trừ khi nó làm sai cấu trúc câu nghiêm trọng. Đây là lý do tại sao một bài có lỗi từ vựng vẫn có thể đạt 9.0 GRA nếu cấu trúc câu phức tạp và đa dạng.
*   **Word Choice:** Ưu tiên "Proportion" cho dữ liệu nhân lực/dân số. "Percentage" chỉ là con số thuần túy.
*   **Precision:** "Chosen one" -> Sai style. Sửa thành "Popular sector".
#### D. Grammatical Range & Accuracy (GRA)
*   **Độ chính xác tuyệt đối:** Soi kỹ từng lỗi mạo từ, giới từ, số ít/nhiều.
*   **Tỷ lệ câu không lỗi (Error-free sentences):**
    *   Band 6: Có lỗi nhưng không quá khó hiểu.
    *   Band 7: Câu không lỗi xuất hiện thường xuyên (Frequent).
    *   Band 8+: Đa số các câu hoàn toàn sạch lỗi (Majority error-free).
*   **Các lỗi kỹ thuật:**
    *   **Comma Splice:** Dùng dấu phẩy nối hai mệnh đề độc lập -> Kéo điểm xuống Band 5-6.
    *   **The Mad Max:** Lạm dụng hoặc thiếu mạo từ "the".
    *   **Past Perfect Trigger:** Thấy "By + [thời gian quá khứ]" mà không dùng Quá khứ hoàn thành -> Đánh dấu yếu kém về Range.
    *   **>> BỔ SUNG QUY TẮC DẤU CÂU (Punctuation Control):** Ngoài Comma Splice, nếu bài viết thường xuyên thiếu dấu phẩy ngăn cách mệnh đề phụ (Subordinate clause), hoặc viết hoa tùy tiện -> **KHÔNG ĐƯỢC CHẤM BAND 8.0 GRA**.
*   **>> CHIẾN THUẬT PARAPHRASING (Introduction Strategy):**
        *   Kiểm tra câu mở đầu (Introduction). Nếu thí sinh chỉ thay từ đồng nghĩa (synonyms) trong cụm danh từ (Noun Phrase), hãy đánh giá ở mức "Standard".
        *   Nếu thí sinh chuyển đổi được cấu trúc từ Noun Phrase (*the number of...*) sang Noun Clause (*how many...*), hãy ghi nhận đây là điểm cộng lớn cho Band 8+ GRA.
*   **Band 9 Threshold:** Nếu bài viết dùng câu phức hay và tự nhiên, cho phép 1-2 lỗi nhỏ (slips). Đừng kẹt ở Band 8.0 chỉ vì một lỗi mạo từ.
*   **>> NGUYÊN TẮC "SLIPS" TRONG GRA:** Band 9.0 GRA cho phép "rare minor errors" (các lỗi nhỏ hiếm gặp). Nếu bài viết sử dụng nhiều cấu trúc phức tạp một cách tự nhiên, đừng ngần ngại cho 9.0 dù vẫn còn 1-2 lỗi mạo từ hoặc số ít/nhiều. Đừng máy móc chặn ở 8.0.
*   **>> GIAO THỨC "PREPOSITION MICRO-SCANNING" (Soi Giới từ Chết người):**
    *   Sau khi quét toàn bộ bài viết, hãy thực hiện một lượt quét **thứ hai** chỉ để tìm lỗi giới từ đi kèm với số liệu và xu hướng.
    *   **To:** Dùng cho điểm đến cuối cùng (VD: "recovered **to** 15%").
    *   **At:** Dùng cho một điểm cố định (VD: "stood **at** 10%").
    *   **Of:** Dùng để chỉ giá trị của một danh từ (VD: "a level **of** 15%").
    *   **In:** Dùng cho năm (VD: "**in** 2015").
    *   **By:** Dùng để chỉ một lượng thay đổi (VD: "decreased **by** 5%").
    *   **BẮT BUỘC:** Nếu thí sinh dùng sai bất kỳ giới từ nào trong các trường hợp trên (ví dụ: dùng "at" hoặc "by" thay vì "to"), hãy bắt lỗi **"Preposition Error"** và giải thích rõ quy tắc sử dụng. Đây là lỗi cơ bản nhưng làm mất điểm rất nặng.
    
### 3. QUY TRÌNH CHẤM ĐIỂM & TỰ SỬA LỖI (SCORING & SELF-CORRECTION)

Mọi từ hoặc dấu câu nằm trong thẻ `<del>...</del>` ở bản sửa **BẮT BUỘC** phải có một mục nhập (entry) riêng biệt tương ứng trong danh sách `errors`. Tuyệt đối không được tóm tắt hay gộp lỗi.
**Bước 1: Deep Scan & Lập danh sách lỗi (JSON Errors Array)**
*   Dựa trên kết quả quét 3 lớp, liệt kê **TẤT CẢ** vấn đề vào mảng `errors`.
*   **>> QUY TẮC "BẰNG CHỨNG BẮT BUỘC" (MANDATORY EVIDENCE):**
    *   Nếu bạn định chấm điểm **Coherence & Cohesion dưới 9.0**, bạn **BẮT BUỘC** phải tạo ra ít nhất **2-3 mục lỗi** trong mảng `errors` thuộc nhóm `Coherence & Cohesion` để giải thích lý do trừ điểm.
    *   *Ví dụ:* Nếu chấm CC 6.0, bạn phải chỉ ra cụ thể: "Đoạn 2 thiếu câu chủ đề", "Từ nối 'Moreover' dùng sai", hoặc "Mạch văn bị đứt gãy".
    *   **CẤM:** Tuyệt đối không được để trống danh sách lỗi CC nếu điểm CC < 9.0.
*   **Thực hiện quét 2 lớp:** 
        *   *Lớp 1 (Grammar/Vocab):* Soi từng mạo từ, dấu phẩy, số ít/nhiều.
        *   *Lớp 2 (Data Logic):* Kiểm tra lỗi "Object vs Figure" (vd: nhầm giữa chủ thể ngành công nghiệp và lượng khí thải). 
*   **Liệt kê toàn bộ lỗi vào mảng `errors` trước.** Nếu có 14 vị trí sai, phải có 14 mục lỗi trong JSON. *Ví dụ:* Nếu sai 3 mạo từ 'the', phải có 3 mục lỗi riêng biệt.
*   **>> QUY TẮC "DOUBLE-TAGGING" (GẮN NHÃN KÉP - MỚI THÊM):**
    *   Nếu gặp lỗi ngữ pháp nghiêm trọng làm đứt gãy mạch văn (như `Sentence Fragment`, `Run-on Sentence`, `Comma Splice`), bạn phải tạo **2 mục lỗi** trong JSON:
        1.  Một mục `Grammar` (để sửa câu chữ).
        2.  Một mục `Coherence & Cohesion` với tên lỗi `Fragmented Flow` (để cảnh báo về mạch lạc).
    *   Điều này đảm bảo phần Coherence & Cohesion không bị trống và không hiển thị thông báo "Tuyệt vời" sai lệch.
*   Dựa trên danh sách lỗi này để tính toán Band điểm cho bài gốc (Markdown).
*   **Quy tắc làm tròn điểm bài viết theo chuẩn IELTS:**
    *   Làm tròn đến nửa band gần nhất (.0 hoặc .5).
    *   **NGOẠI LỆ BẮT BUỘC:**
        *   Điểm trung bình có đuôi **.25** -> BẮT BUỘC làm tròn **XUỐNG** số nguyên (Ví dụ: 8.25 -> 8.0).
        *   Điểm trung bình có đuôi **.75** -> BẮT BUỘC làm tròn **XUỐNG** .5 (Ví dụ: 8.75 -> 8.5).

**Bước 2: Tạo bản sửa lỗi (Annotated Essay)**
    *   **Nguyên tắc "Soi gương":** Bạn chỉ được phép sửa lỗi dựa trên danh sách lỗi đã lập ở Bước 1. 
    *   **Cấm sửa ngầm (No Hidden Edits):** Tuyệt đối không được "tiện tay" sửa các lỗi nhỏ (như thêm mạo từ 'the' hay viết hoa) trong bài sửa nếu bạn chưa khai báo lỗi đó trong danh sách `errors` ở Bước 1. 
    *   **Số lượng thẻ `<del>` phải bằng chính xác số lượng lỗi trong JSON.** Nếu sai lệch, hệ thống sẽ coi là vi phạm giao thức.
    
**Bước 3: Chấm lại bản sửa lỗi (JSON Output - Internal Re-grading)**
*   Hãy đóng vai một Giám khảo độc lập thứ 2 chấm lại bản `annotated_essay` vừa tạo (coi đây là một bài nộp mới đã sạch lỗi câu chữ).
*   **Luật Nội dung (Content Rule):** Vì bản sửa này chỉ khắc phục GRA/LR và giữ nguyên cấu trúc cũ, nên điểm TA và CC của bản sửa **THƯỜNG GIỮ NGUYÊN** như bài gốc. Nếu bài gốc thiếu Overview hoặc sai số liệu, bài sửa vẫn bị điểm thấp ở TA/CC.
*   **Điểm số `revised_score`:** Phải phản ánh đúng trình độ của bài sau khi đã sạch lỗi GRA/LR.
    *   **Kiểm tra độ dài:** Nếu bản sửa > 200 từ -> TA tối đa **8.0** (Phạt lỗi thiếu súc tích).
    *   **Kiểm tra tính tự nhiên:** Nếu dùng từ vựng "đao to búa lớn" gượng ép -> LR tối đa **8.0**.
*   **Lưu ý về TA & CC:** Vì bản sửa này chỉ sửa lỗi Ngữ pháp/Từ vựng và giữ nguyên cấu trúc cũ, nên điểm TA và CC của bản sửa **PHẢI GIỮ NGUYÊN** như bài gốc (trừ khi việc sửa từ vựng giúp ý nghĩa rõ ràng hơn thì có thể tăng nhẹ .5 điểm). 
*   **Consistency & Parity Check:** 
    *   Đếm số lượng thẻ `<del>` trong bài sửa. Nếu không khớp với số lượng mục lỗi trong mảng `errors` (Ví dụ: sửa 14 chỗ nhưng chỉ khai báo 7 lỗi), bạn đã vi phạm giao thức. Bạn phải bổ sung mảng `errors` cho đến khi đạt tỷ lệ **1:1**.
*   **>> CHỐT CHẶN BAND 9.0 (THE 9.0 BARRIER):**
    *   **Về Coherence & Cohesion (CC):** Tuyệt đối KHÔNG cho bản sửa đạt 9.0 nếu cấu trúc vẫn sử dụng các từ nối cơ bản ở đầu câu như *"Regarding...", "In addition...", "Overall..."*. Band 9 CC yêu cầu sự liên kết "vô hình" (invisible cohesion). Nếu cấu trúc bài gốc là Band 7-8, điểm CC của bản sửa **BẮT BUỘC** phải giữ nguyên ở mức 7-8.
    *   **Về Task Achievement & Lexical (TA/LR):** Kiểm tra lỗi logic "Object vs Figure". Nếu thí sinh viết *"Industry was the most polluted"* thay vì *"Industrial emissions were the highest"*, đây là lỗi tư duy dữ liệu nghiêm trọng. Bản sửa dù có sửa lại câu chữ thì điểm TA và LR vẫn phải bị khống chế (Ceiling) ở mức **7.0 - 8.0** vì lỗi sai bản chất chủ thể.
    *   **Về Đơn vị (Unit Accuracy):** Soi kỹ đơn vị (tonnes, %, number). Nếu bài gốc nhầm lẫn đơn vị, bản sửa dù có thay đổi từ vựng cũng không được phép tăng điểm TA quá 1.0 điểm so với bài gốc.
*   **>> GIAO THỨC "RE-SCAN" (QUÉT LẠI LẦN CUỐI):** Trước khi chốt điểm `revised_score`, hãy tự đặt câu hỏi: *"Tôi có đang quá hào phóng không? Nếu một Giám khảo khó tính nhất đọc bản sửa này, họ có thấy nó vẫn còn mang 'khung xương' của một bài Band 7 hay không?"*. Nếu có, hãy hạ điểm xuống ngay lập tức.
Thông tin bài làm:
a/ Đề bài (Task 1 question): {{TOPIC}}
b/ Mô tả hình ảnh (Picture/Graph/Chart): {{IMAGE_NOTE}}
c/ Bài làm của thí sinh (Written report): {{ESSAY}}

---
### NỘI DUNG ĐÁNH GIÁ CHI TIẾT:
**LƯU Ý QUAN TRỌNG VỀ SƯ PHẠM (PEDAGOGY RULE):**
Khi đưa ra ví dụ sửa lỗi (Example/Rewrite), bạn phải căn cứ vào **Band điểm hiện tại** của bài làm:
*   **Nếu bài < 6.0:** Hãy đưa ra ví dụ sửa ở mức **Band 7.0** (Tập trung vào sự Chính xác, Rõ ràng, Dễ hiểu). Đừng dùng từ quá khó.
*   **Nếu bài >= 6.5:** Hãy đưa ra ví dụ sửa ở mức **Band 9.0** (Tập trung vào sự Tinh tế, Học thuật, Cấu trúc phức tạp).
**QUY TẮC "CHỐNG SƠ SÀI" (ANTI-BREVITY RULE):**
1.  **Cấm nhận xét chung chung:** Tuyệt đối không viết "Cần cải thiện ngữ pháp" mà không chỉ rõ là cải thiện cái gì (thì, mạo từ, hay cấu trúc?).
2.  **Trích dẫn bằng chứng:** Mọi nhận xét đều phải trích dẫn câu văn cụ thể của thí sinh để chứng minh.
3.  **Luôn viết mẫu:** Dù bài làm ở Band 1 hay Band 9, bạn **BẮT BUỘC** phải cung cấp các ví dụ viết lại (Rewrite) ở cuối mỗi tiêu chí. Không được bỏ qua.

### **1. Task Achievement (Hoàn thành yêu cầu bài thi):**
*   **Kiểm tra Introduction (Mở bài):**
    *   [Xác định xem bài viết có câu mở đầu không? Thí sinh đã paraphrase đề bài bằng cách nào (Thay từ hay Đổi cấu trúc)?]
    *   **⚠️ Cảnh báo:** [Nếu thiếu Introduction, hãy tuyên bố ngay lập tức: "Bạn đã vi phạm lỗi Format nghiêm trọng. Điểm TA của bạn bị giới hạn ở Band 5.0 bất kể thân bài viết hay đến đâu."]
    *   **⚠️ Cảnh báo sao chép:** [Nếu chép đề: "Bạn đang sao chép lại đề bài. Những từ ngữ này sẽ không được tính điểm."]
*   **Đánh giá Overview (Cái nhìn tổng quan):** 
    *   [Phân tích: Đã có Overview chưa? Có nêu được xu hướng chính và sự so sánh nổi bật không?]
    *   **⚠️ Cảnh báo cho trình độ Band 5-6:** [Nếu Overview vẫn bị dính số liệu chi tiết, hãy giải thích tại sao lỗi này khiến họ bị kẹt ở Band 5 và hướng dẫn cách xóa bỏ để lên Band 7.]
*   **Độ chính xác và Chọn lọc dữ liệu:** 
    *   [Kiểm tra độ chính xác của số liệu. Có bị lỗi "Data Saturation" - nhồi nhét quá nhiều số liệu vụn vặt không?]
    *   [**Lưu ý:** Bỏ qua dữ liệu 'Total'/'Other' nếu không quan trọng.]
*   **Giải quyết yêu cầu (Response Strategy):** [Đánh giá cách nhóm thông tin. Thí sinh đang mô tả đơn lẻ (Band 5) hay đã biết tổng hợp dữ liệu để so sánh (Band 7+)?]

*   **⚠️ Các lỗi nghiêm trọng & Phân tích chuyên sâu:** 
    *   [Với mỗi lỗi tìm được, bạn **BẮT BUỘC** giải thích theo 3 bước:
        1. **Trích dẫn lỗi:** (Ví dụ: "the figure of pizza ate")
        2. **Lý do yếu kém:** (Ví dụ: Vi phạm lỗi tư duy Object vs Figure).
        3. **Tác động:** (Ví dụ: Làm mất tính chuyên nghiệp, khiến giám khảo đánh giá thấp tư duy logic).]

*   **💡 CHIẾN THUẬT NÂNG BAND (STEP-BY-STEP):**
    *   **Bước 1 (Lọc):** Tuyệt đối xóa số liệu khỏi Overview. Overview chỉ nói về "ý nghĩa" con số.
    *   **Bước 2 (Gộp):** Nhóm các đối tượng cùng tăng/cùng giảm để tạo sự súc tích (Economy).
    *   **Bước 3 (So sánh):** Luôn phải chỉ ra điểm cao nhất/thấp nhất hoặc sự thay đổi thứ hạng đáng kể.
    *   **Bước 4 (Kết nối):** Sử dụng liên kết "tàng hình" (While/Whereas/V-ing) thay vì từ nối máy móc.
    
*   **✍️ HÌNH MẪU ĐỐI CHIẾU (CHỌN MỨC PHÙ HỢP ĐỂ HỌC):**
    *   **Mẫu thực tế (Mục tiêu Band 7.0):** 
        *   *"Đây là phiên bản rõ ràng, chính xác, không lỗi logic mà bạn có thể đạt được ngay sau khi chỉnh sửa bài làm hiện tại:"*
        *   **[AI HÃY VIẾT OVERVIEW & BODY ĐẠT CHUẨN 7.0 DỰA TRÊN Ý TƯỞNG CỦA HỌC VIÊN]**
    *   **Mẫu chuyên sâu (Tham khảo Band 9.0):** 
        *   *"Đây là phiên bản để bạn tham khảo cách dùng từ vựng tinh tế và cấu trúc tổng hợp dữ liệu đỉnh cao của Giám khảo:"*
        *   **[AI HÃY VIẾT OVERVIEW & BODY ĐẠT CHUẨN 9.0 TẠI ĐÂY]**

> **📍 Điểm Task Achievement:** [Điểm số/9.0]

#### **2. Coherence and Cohesion (Độ mạch lạc và liên kết):**

*   **Tổ chức đoạn văn (Paragraphing):** [Phân tích logic chia đoạn: Bạn chia đoạn theo Tiêu chí gì (Thời gian/Đối tượng/Xu hướng)? Cách chia này có giúp người đọc dễ so sánh không? Mỗi đoạn có một trọng tâm rõ ràng không?]
*   **Sử dụng từ nối (Linking Devices):** [Đánh giá độ tự nhiên:
    *   **Cảnh báo:** Có bị lạm dụng từ nối đầu câu ("Mechanical Linking") như *Regarding, Turning to, Looking at, Firstly* không?
    *   **Khuyến khích:** Có sử dụng "Invisible Cohesion" (trạng từ đứng giữa câu như *meanwhile, however* hoặc dùng mệnh đề quan hệ để nối ý) không?]
*   **Phép tham chiếu (Referencing):** [Kiểm tra kỹ thuật Referencing: Bạn có sử dụng *it, this, that, the former, the latter, respectively* để tránh lặp từ không? Hay bạn lặp lại danh từ liên tục?]
*   **⚠️ Lỗi cần khắc phục:** [Chỉ ra cụ thể (càng nhiều càng tốt):
    1.  **Mạch văn đứt gãy:** Các câu rời rạc, không ăn nhập.
    2.  **Tham chiếu sai:** Dùng "it" nhưng không rõ thay thế cho từ nào (Ambiguous Reference).
    3.  **Lỗi cấu trúc:** Lặp lại cấu trúc câu (VD: Câu nào cũng bắt đầu bằng "The figure...").
    4.  **Câu thiếu động từ (Fragment):** Gây khó hiểu.]
*   **💡 Cải thiện & Nâng cấp (Correction & Upgrade):**
    *   *Câu gốc (Vấn đề):* "[Trích dẫn chính xác câu văn bị máy móc/lủng củng của thí sinh]"
    *   *Gợi ý viết lại (Natural Flow):* "[Nếu Band thấp: Sửa cho ĐÚNG ngữ pháp và RÕ nghĩa nối. Nếu Band 7+: Viết lại câu đó sử dụng cấu trúc liên kết ẩn hoặc chủ ngữ liên kết để đạt Band 8-9]"
    *   *Giải thích:* "[Tại sao cách viết mới giúp bài văn mượt mà và chuyên nghiệp hơn?]"
* **Yêu cầu bắt buộc về độ sâu:** Với mỗi lỗi tìm được, bạn phải giải thích theo 3 bước:
    1. Trích dẫn lỗi.
    2. Giải thích tại sao quy tắc Band Descriptors coi đây là điểm yếu.
    3. Phân tích tác động của lỗi này đến người đọc (gây hiểu lầm, làm mất tính chuyên nghiệp...).
    
> **📍 Điểm Coherence & Cohesion:** [Điểm số/9.0]

#### **3. Lexical Resource (Vốn từ vựng):**

*   **Đánh giá độ đa dạng (Range & Flexibility):** [Nhận xét tổng quan: Vốn từ của thí sinh đang ở mức nào? (Cơ bản/Đủ dùng/Phong phú). Có bị lỗi lặp từ ("Repetition") nghiêm trọng với các từ khóa chính (increase, decrease, figure...) không?]
*   **Độ chính xác và Văn phong (Precision & Style):** [Đánh giá: Thí sinh có dùng được các cụm từ kết hợp (Collocations) tự nhiên không hay là dịch từ tiếng mẹ đẻ (Word-for-word translation)? Có từ nào bị dùng sai ngữ cảnh (ví dụ: dùng văn nói "get up" thay vì "increase") không?]
*   **⚠️ Điểm yếu cốt lõi:** [Đừng liệt kê từng lỗi chính tả. Hãy chỉ ra **thói quen sai** của thí sinh. Ví dụ: *"Bạn thường xuyên chọn sai từ để mô tả đối tượng (Object)"* hoặc *"Bạn lạm dụng từ vựng quá trang trọng (Pretentious) không cần thiết"*.]
*   **💡 Gợi ý nâng cấp (Vocabulary Upgrade):**
    *   *Thay thế từ vựng thường:* "[Tìm 1 từ lặp lại nhiều nhất trong bài, ví dụ 'increase']"
    *   *Gợi ý thay thế:* 
        *   *[Nếu Band < 7]:* Gợi ý các từ cơ bản nhưng đúng (rise, growth, go up).
        *   *[Nếu Band 7+]:* Gợi ý các từ học thuật (escalate, upsurge, register a growth).
* **Yêu cầu bắt buộc về độ sâu:** Với mỗi lỗi tìm được, bạn phải giải thích theo 3 bước:
    1. Trích dẫn lỗi.
    2. Giải thích tại sao quy tắc Band Descriptors coi đây là điểm yếu.
    3. Phân tích tác động của lỗi này đến người đọc (gây hiểu lầm, làm mất tính chuyên nghiệp...).
    
> **📍 Điểm Lexical Resource:** [Điểm số/9.0]

#### **4. Grammatical Range and Accuracy (Ngữ pháp):**

*   **Độ đa dạng cấu trúc (Range Check):** [Phân tích chiến lược: Bài viết có "nghèo nàn" cấu trúc không? (Chỉ dùng câu đơn/câu ghép cơ bản). Thí sinh có sử dụng được các cấu trúc Band 8+ không: *Passive Voice (Bị động)*, *Reduced Relative Clause (Rút gọn mệnh đề)*, *Nominalization (Danh từ hóa)*?]
*   **Độ chính xác (Accuracy Check):** [Ước lượng tỷ lệ câu không lỗi (Error-free sentences): Dưới 50% (Band 5), 50-70% (Band 6-7), hay trên 80% (Band 8+)? Lỗi sai chủ yếu là lỗi hệ thống (Systematic - sai quy tắc) hay lỗi sơ suất (Slips)?].Nếu bài viết có trên 80% số câu hoàn toàn sạch lỗi (Error-free) và lỗi duy nhất là một lỗi nhỏ (như "most highest") -> **Vẫn giữ mức Band 8.5 - 9.0**. Đừng ép thí sinh dùng cấu trúc lạ nếu cấu trúc hiện tại đã quá đủ để truyền đạt thông tin một cách tinh tế. Band 9 không bắt buộc phải có "Đảo ngữ" hay "Câu điều kiện". Range được thể hiện qua việc sử dụng linh hoạt: Mệnh đề quan hệ, câu phân từ (Reduced clauses), danh từ hóa (Nominalization), và các cấu trúc so sánh phức tạp. 
*   **Dấu câu (Punctuation):** [Nhận xét việc dùng dấu phẩy, dấu chấm. Có mắc lỗi *Comma Splice* (Dấu phẩy nối câu) kinh điển không?]
*   **⚠️ Lỗi hệ thống cần sửa:** [Chỉ ra lỗ hổng kiến thức ngữ pháp lớn nhất của thí sinh. Ví dụ: *"Bạn rất yếu về Mệnh đề quan hệ"* hoặc *"Bạn chưa nắm vững cách dùng Mạo từ"*.]
*   **💡 Thử thách viết lại (Sentence Transformation):**
    *   *Câu gốc (Simple/Error):* "[Trích 1 câu đơn giản hoặc có lỗi trong bài]"
    *   *Nâng cấp câu:* 
        *   *[Nếu Band thấp]:* Ghép thành câu ghép/câu phức cơ bản (dùng because, although) để đảm bảo đúng.
        *   *[Nếu Band cao]:* Dùng cấu trúc nâng cao (Mệnh đề phân từ, Đảo ngữ, Nominalization).
* **Yêu cầu bắt buộc về độ sâu:** Với mỗi lỗi tìm được, bạn phải giải thích theo 3 bước:
    1. Trích dẫn lỗi.
    2. Giải thích tại sao quy tắc Band Descriptors coi đây là điểm yếu.
    3. Phân tích tác động của lỗi này đến người đọc (gây hiểu lầm, làm mất tính chuyên nghiệp...).
    
> **📍 Điểm Grammatical Range & Accuracy:** [Điểm số/9.0]

---
### **TỔNG ĐIỂM (OVERALL BAND SCORE):** Quy tắc làm tròn điểm bài viết theo chuẩn IELTS:
    *   Làm tròn đến nửa band gần nhất (.0 hoặc .5).
    *   **NGOẠI LỆ BẮT BUỘC:**
        *   Điểm trung bình có đuôi **.25** -> BẮT BUỘC làm tròn **XUỐNG** số nguyên (Ví dụ: 8.25 -> 8.0).
        *   Điểm trung bình có đuôi **.75** -> BẮT BUỘC làm tròn **XUỐNG** .5 (Ví dụ: 8.75 -> 8.5).

---
### **LỜI KHUYÊN CHIẾN THUẬT TỪ GIÁM KHẢO (EXAMINER'S TIPS):**
1.  **Đưa ra các lời khuyên:** Hãy đưa ra các lời khuyên chiến thuật dựa trên những lỗi sai thực tế trong bài.
2.  **Economy:** Cách cắt giảm số từ thừa (nếu bài > 200 từ).
3.  **Introduction Power:** Cách đổi Noun Phrase -> Noun Clause trong mở bài.
4.  **Grouping:** Cách nhóm thông tin thông minh hơn (nhóm theo xu hướng Lớn vs Nhỏ).
5.  **Overview:** Cách viết Overview tốt hơn.

#### **5. DỮ LIỆU PHÂN TÍCH (ANALYSIS DATA):**

Sau khi đánh giá xong, bạn **BẮT BUỘC** phải trích xuất dữ liệu dưới dạng một **JSON Object duy nhất**.

**QUAN TRỌNG:** Trong trường "type" (Tên lỗi), bạn CHỈ ĐƯỢC PHÉP được dùng các thuật ngữ tiếng Anh chuẩn học thuật dưới đây:

**A. [COHERENCE & COHESION] - Macro Errors:**
# Organization & Progression (Tổ chức & Phát triển)
`Illogical Grouping` (Sắp xếp phi logic), `Missing Overview` (Thiếu tổng quan), `Fragmented Flow` (Mạch văn đứt gãy), `Lack of Progression` (Không phát triển ý), `Incoherent Paragraphing` (Chia đoạn không mạch lạc).
# Linking & Reference (Liên kết & Tham chiếu)
`Mechanical Linking` (Từ nối máy móc), `Overuse of Connectors` (Lạm dụng từ nối), `Ambiguous Referencing` (Tham chiếu mơ hồ), `Repetitive Structure` (Lặp cấu trúc), `Data Inaccuracy` (Sai số liệu/Logic).

**B. [GRAMMAR] - Micro Errors:**
# Sentence Structure (Cấu trúc câu)
`Comma Splice` (Lỗi dấu phẩy), `Run-on Sentence` (Câu dính liền), `Sentence Fragment` (Câu thiếu thành phần), `Faulty Parallelism` (Lỗi song song), `Misplaced Modifier` (Bổ ngữ sai chỗ), `Word Order` (Trật tự từ).
# Morphology & Syntax (Hình thái & Cú pháp)
`Subject-Verb Agreement` (Hòa hợp chủ vị), `Tense Inconsistency` (Sai thì), `Passive Voice Error` (Lỗi bị động), `Relative Clause Error` (Lỗi mệnh đề quan hệ).
# Mechanics (Cơ học)
`Article Error` (Mạo từ), `Preposition Error` (Giới từ), `Singular/Plural` (Số ít/nhiều), `Countable/Uncountable` (Danh từ đếm được/không), `Punctuation` (Dấu câu).

**C. [VOCABULARY] - Lexical Errors:**
# Meaning & Use (Nghĩa & Cách dùng)
`Imprecise Word Choice` (Dùng từ thiếu chính xác), `Incompatible Collocation` (Kết hợp từ sai), `Word Form Error` (Sai loại từ), `Selectional Restriction Violation` (Vi phạm quy tắc chọn lọc từ).
# Style & Register (Văn phong)
`Informal Register` (Văn phong suồng sã), `Pretentious Language` (Dùng từ sáo rỗng/làm màu), `Redundancy` (Thừa từ/Lặp ý), `Forced Paraphrasing` (Paraphrase gượng ép).

**CATEGORY MAPPING RULE:**
*   Group A -> `category`: "Coherence & Cohesion"
*   Group B -> `category`: "Grammar"
*   Group C -> `category`: "Vocabulary"

**TỰ CHẤM LẠI BẢN SỬA (INTERNAL RE-GRADING - BƯỚC QUAN TRỌNG NHẤT):**
   - Hãy quên rằng bạn vừa sửa bài này. Hãy đóng vai một Giám khảo độc lập thứ 2 chấm lại bản 'annotated_essay' vừa tạo.
   - **Luật Nội dung (Content Rule):** Bản sửa chỉ sửa ngữ pháp/từ vựng, KHÔNG THỂ sửa lỗi thiếu số liệu/thiếu so sánh của bài gốc. Nếu bài gốc TA 6.0, bản sửa TA vẫn là 6.0 (hoặc tối đa 7.0 nếu diễn đạt rõ hơn).
   - **Kết luận:** Điểm 'revised_score' PHẢI là điểm thực tế của bản sửa, KHÔNG ĐƯỢC mặc định là 9.0.
Cấu trúc JSON:
```json
{
  "original_score": {
      "task_achievement": "Điểm TA của bài làm gốc (User's essay)",
      "cohesion_coherence": "Điểm CC của bài làm gốc",
      "lexical_resource": "Điểm LR của bài làm gốc",
      "grammatical_range": "Điểm GRA của bài làm gốc",
      "overall": "Điểm Overall của bài làm gốc (Average)"
  },
  "errors": [
    {
      "category": "Grammar" hoặc "Vocabulary",
      "type": "Tên Lỗi",
      "impact_level": "High" | "Medium" | "Low",
      "explanation": "Giải thích ngắn gọn lỗi.",
      "original": "đoạn văn bản sai",
      "correction": "đoạn văn bản đúng (VIẾT IN HOA)"
    }
  ],
  "annotated_essay": "Phiên bản bài làm đã được sửa lỗi (giữ nguyên cấu trúc các đoạn văn). Bọc từ sai trong thẻ <del>...</del> và từ sửa đúng trong thẻ <ins class='grammar'>...</ins> hoặc <ins class='vocab'>...</ins>. Nội dung sửa đúng phải viết IN HOA.",
   "revised_score": {
      "word_count_check": "BẮT BUỘC GHI SỐ TỪ CỦA BẢN SỬA (Ví dụ: '220 words - Too long')",
      "logic_re_evaluation": "Giải thích tại sao bị trừ điểm (Ví dụ: 'Dù sạch lỗi ngữ pháp nhưng bài viết dài 220 từ, vi phạm nguyên tắc súc tích, nên TA chỉ đạt 8.0').",
      "task_achievement": "Điểm TA thực tế (phạt nặng nếu dài dòng)",
      "cohesion_coherence": "Điểm CC",
      "lexical_resource": "Điểm LR",
      "grammatical_range": "Điểm GRA",
      "overall": "Điểm trung bình (Làm tròn theo Quy tắc làm tròn điểm bài viết theo chuẩn IELTS)"
          *   Làm tròn đến nửa band gần nhất (.0 hoặc .5).
          *   **NGOẠI LỆ BẮT BUỘC:**
              *   Điểm trung bình có đuôi **.25** -> BẮT BUỘC làm tròn **XUỐNG** số nguyên (Ví dụ: 8.25 -> 8.0).
              *   Điểm trung bình có đuôi **.75** -> BẮT BUỘC làm tròn **XUỐNG** .5 (Ví dụ: 8.75 -> 8.5).
  }
}
```
"""

# ==========================================
# 3. HELPER FUNCTIONS
# ==========================================

def clean_json(text):
    # Tìm đoạn văn bản nằm giữa dấu ngoặc nhọn { ... } đầu tiên và cuối cùng
    match = re.search(r"(\{[\s\S]*\})", text)
    if match:
        content = match.group(1).strip()
        # Loại bỏ các ký tự điều khiển lỗi
        content = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', content)
        return content
    return None

def parse_guide_response(text):
    j_str = clean_json(text)
    if not j_str: return None
    try:
        return json.loads(j_str)
    except:
        # Nếu lỗi JSON, thử quét tay các trường quan trọng (fallback)
        return {
            "task_type": "IELTS Task 1",
            "intro_guide": "Hãy paraphrase đề bài bằng từ đồng nghĩa.",
            "overview_guide": "Nêu xu hướng chung và đặc điểm nổi bật.",
            "body1_guide": "Mô tả chi tiết nhóm số liệu 1.",
            "body2_guide": "Mô tả chi tiết nhóm số liệu 2."
        }

def process_grading_response(full_text):
    """
    Hàm xử lý kết quả chấm điểm (CHUẨN TỪ APP CHẤM ĐIỂM).
    Tách biệt:
    1. Markdown Text (Phân tích chi tiết ở đầu).
    2. JSON Data (Điểm số và lỗi ở cuối).
    """
    json_str = clean_json(full_text)
    
    # Mặc định
    markdown_part = full_text
    data = {
        "errors": [], 
        "annotatedEssay": None, 
        "revisedScore": None, 
        "originalScore": {
            "task_achievement": "-", "cohesion_coherence": "-", 
            "lexical_resource": "-", "grammatical_range": "-", "overall": "-"
        }
    }
    
    if json_str:
        # Tách phần Markdown (trước JSON)
        markdown_part = full_text.split("```json")[0].strip()
        # Nếu AI không dùng code block, thử split bằng ký tự '{' đầu tiên của JSON
        if "original_score" in markdown_part: # Dấu hiệu JSON bị lẫn
             parts = full_text.split("{", 1)
             markdown_part = parts[0].strip()

        try:
            parsed = json.loads(json_str)
            data["errors"] = parsed.get("errors", [])
            data["annotatedEssay"] = parsed.get("annotated_essay")
            data["revisedScore"] = parsed.get("revised_score")
            data["originalScore"] = parsed.get("original_score", {})
        except json.JSONDecodeError:
            pass

    return markdown_part, data

# --- FILE EXPORT ---
def register_vietnamese_font():
    try:
        font_reg = "Roboto-Regular.ttf"
        font_bold = "Roboto-Bold.ttf"
        if not os.path.exists(font_reg):
            r = requests.get("https://github.com/googlefonts/roboto/raw/main/src/hinted/Roboto-Regular.ttf")
            with open(font_reg, "wb") as f: f.write(r.content)
        if not os.path.exists(font_bold):
            r = requests.get("https://github.com/googlefonts/roboto/raw/main/src/hinted/Roboto-Bold.ttf")
            with open(font_bold, "wb") as f: f.write(r.content)
        pdfmetrics.registerFont(TTFont('Roboto', font_reg))
        pdfmetrics.registerFont(TTFont('Roboto-Bold', font_bold))
        addMapping('Roboto', 0, 0, 'Roboto')
        addMapping('Roboto', 1, 0, 'Roboto-Bold')
        return True
    except: return False

def create_docx(data, topic, essay, analysis):
    doc = Document()
    doc.add_heading('IELTS ASSESSMENT REPORT', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('1. DETAILED ANALYSIS', level=1)
    doc.add_paragraph(analysis) # Phân tích chi tiết từ Markdown
    
    # Thêm bảng điểm
    doc.add_heading('2. SCORE BREAKDOWN', level=1)
    scores = data.get("originalScore", {})
    p = doc.add_paragraph()
    p.add_run(f"Overall Band: {scores.get('overall', '-')}\n").bold = True
    p.add_run(f"TA: {scores.get('task_achievement', '-')}, CC: {scores.get('cohesion_coherence', '-')}, LR: {scores.get('lexical_resource', '-')}, GRA: {scores.get('grammatical_range', '-')}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf(data, topic, essay, analysis):
    register_vietnamese_font()
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = [Paragraph("IELTS ASSESSMENT REPORT", styles['Title'])]
    
    # Analysis
    elements.append(Paragraph("DETAILED ANALYSIS", styles['Heading1']))
    # Clean markdown basic symbols for PDF
    safe_text = html.escape(analysis).replace('\n', '<br/>').replace('**', '').replace('#', '')
    elements.append(Paragraph(safe_text, styles['Normal']))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer

# ==========================================
# 4. UI: QUẢN LÝ TRẠNG THÁI (SESSION STATE)
# ==========================================
if "step" not in st.session_state: st.session_state.step = 1 
if "guide_data" not in st.session_state: st.session_state.guide_data = None
if "grading_result" not in st.session_state: st.session_state.grading_result = None
if "saved_topic" not in st.session_state: st.session_state.saved_topic = ""
if "saved_img" not in st.session_state: st.session_state.saved_img = None

# ==========================================
# 5. GIAO DIỆN CHÍNH (THEO YÊU CẦU MỚI)
# ==========================================

# TIÊU ĐỀ CHÍNH
st.markdown('<div class="main-header">🎓 IELTS Writing Task 1 – Examiner-Guided</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Learning & Scoring Based on IELTS Band Descriptors</div>', unsafe_allow_html=True)

if st.session_state.step == 1:
    
    # STEP 1 – Task 1 Question (Đã đổi lên trên)
    st.markdown('<div class="step-header">STEP 1 – Task 1 Question</div>', unsafe_allow_html=True)
    st.markdown('<div class="step-desc">Paste the official task question here</div>', unsafe_allow_html=True)
    question_input = st.text_area("Question", height=150, placeholder="The chart below shows...", key="q_input", label_visibility="collapsed")

    # STEP 2 – Visual Data (Đã đổi xuống dưới)
    st.markdown('<div class="step-header">STEP 2 – Visual Data </div>', unsafe_allow_html=True)
    st.markdown('<div class="step-desc">Upload chart / graph / table / diagram / map </div>', unsafe_allow_html=True)
    uploaded_image = st.file_uploader("Upload Image", type=['png', 'jpg', 'jpeg'], key="img_input", label_visibility="collapsed")
    
    img_data = None
    if uploaded_image:
        img_data = Image.open(uploaded_image)
        st.image(img_data, caption='Uploaded Visual Data', width=400)

    # STEP 3    
    st.markdown('<div class="step-header">STEP 3 – Examiner Workflow</div>', unsafe_allow_html=True)
    
    # --- PHẦN HTML NÀY PHẢI VIẾT SÁT LỀ TRÁI (KHÔNG THỤT DÒNG) ---
    workflow_html = """
<style>
    .wf-container { display: grid; grid-template-columns: 1fr 1fr; gap: 15px; margin-bottom: 20px; }
    .wf-card { background-color: #FFFFFF; border: 1px solid #E2E8F0; border-radius: 8px; padding: 15px; display: flex; align-items: center; }
    .wf-icon { width: 40px; height: 40px; background-color: #F0F9FF; color: #0284C7; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-size: 20px; margin-right: 15px; flex-shrink: 0; }
    .wf-title { font-weight: 700; font-size: 0.95rem; color: #1E293B; }
    .wf-desc { font-size: 0.85rem; color: #64748B; line-height: 1.4; }
</style>
<div class="wf-container">
    <div class="wf-card">
        <div class="wf-icon">🔍</div>
        <div class="wf-content">
            <div class="wf-title">1. Task Analysis</div>
            <div class="wf-desc">Analyze visual data to identify chart type.</div>
        </div>
    </div>
    <div class="wf-card">
        <div class="wf-icon">🧠</div>
        <div class="wf-content">
            <div class="wf-title">2. Data Grouping & Planning</div>
            <div class="wf-desc">Organise key features and trends logically..</div>
        </div>
    </div>
    <div class="wf-card">
        <div class="wf-icon">✍️</div>
        <div class="wf-content">
            <div class="wf-title">3. Guided Writing</div>
            <div class="wf-desc">Support writing with clear structure and useful vocabulary.</div>
        </div>
    </div>
    <div class="wf-card">
        <div class="wf-icon">⚖️</div>
        <div class="wf-content">
            <div class="wf-title">4. Band Score Evaluation</div>
            <div class="wf-desc">Evaluate based on official Band Descriptors.</div>
        </div>
    </div>
</div>
"""
    # GỌI LỆNH RENDER
    st.markdown(workflow_html, unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)

    # Nút bấm xử lý (vẫn sử dụng question_input và img_data đã khai báo ở trên)
    if st.button("🔍 Analyze & Guide (Start Learning)", type="primary", width="stretch"):
        if not question_input or not img_data:
            st.warning("⚠️ Vui lòng nhập đầy đủ Đề bài và tải Ảnh lên để bắt đầu.")
        else:
            st.session_state.saved_topic = question_input
            st.session_state.saved_img = img_data
                   
            with st.spinner("🧠 The examiner is analysing the visual data and providing step-by-step guidance on how to write the answer..."):
                    # Prompt Tutor Vạn Năng: Tự động thích ứng theo từng dạng bài
                    prompt_guide = """
                    Bạn là một Siêu Giáo viên IELTS Writing (Band 9.0). Nhiệm vụ của bạn là phân tích hình ảnh đầu vào và viết hướng dẫn thực hành chi tiết.
                    **YÊU CẦU ĐẶC BIỆT (CHẾ ĐỘ PHÂN TÍCH KỸ):** Bạn không cần phải trả lời nhanh. Hãy dành thời gian "suy nghĩ" để phân tích thật sâu và chi tiết (Step-by-step Analysis).
                    # STRICT OUTPUT RULES (BẮT BUỘC TUÂN THỦ):
                    1.  **NO MARKDOWN LISTS:** Tuyệt đối KHÔNG được tự ý chuyển đổi định dạng sang gạch đầu dòng (bullet points) của Markdown.
                    2.  **HTML ONLY:** Output bắt buộc phải giữ nguyên các thẻ HTML: `<ul>`, `<li>`, `<b>`, `<br>`, `<code>`, `<div>`. Hệ thống chỉ render được HTML, nếu bạn dùng Markdown sẽ bị lỗi hiển thị.
                        
                    **BƯỚC 1: NHẬN DIỆN LOẠI BÀI (QUAN TRỌNG)**
                    Hãy nhìn hình ảnh và xác định nó thuộc loại nào:
                    1. **Change Over Time** (Line, Bar, Table, Pie có năm tháng): Cần từ vựng xu hướng (increase, decrease).
                    2. **Static Chart** (Pie, Table, Table 1 năm): Cần từ vựng so sánh (higher, lower, accounts for).
                    3. **Map (Bản đồ):** Cần từ vựng phương hướng (North, South) và sự thay đổi (demolished, constructed). Tuyệt đối không dùng "increase/decrease" cho nhà cửa.
                    4. **Process (Quy trình):** Cần câu Bị động (Passive voice) và từ nối trình tự (First, Then, Finally).
                    5. **Mixed (Kết hợp):** Cần hướng dẫn cách liên kết 2 biểu đồ.
                    
                    
                    **BƯỚC 2: SOẠN HƯỚNG DẪN (OUTPUT JSON)**

                    # =================================================================
                    # 🔴 TRƯỜNG HỢP 1: DẠNG "STATIC COMPARISON" (Pie, Bar, Table - 1 Năm/Không năm)
                    # (Tư duy cốt lõi: Ranking (Xếp hạng) & Proportion (Tỷ trọng))
                    # =================================================================

                    <br><i>(LƯU Ý: Nhìn Năm trong đề bài để quyết định THÌ cho toàn bài viết)</i>
                    <br>Tuyệt đối <b>KHÔNG</b> dùng dấu gạch ngang để chỉ khoảng số (VD: <i>7-14%</i>).
                    <br>👉 <b>Phải viết chữ:</b> <i>"between 7% and 14%"</i> hoặc <i>"from 7% to 14%"</i>.
                    ### 1. **"intro_guide" (Quy trình 5 bước sản xuất Mở bài):**

<ul>
    <!-- ================================================================================== -->
    <!-- PHẦN 1: CÔNG THỨC & LƯU Ý (THIẾT LẬP NỀN TẢNG) -->
    <!-- ================================================================================== -->
    <li>
        <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">             
            <b>► Logic cốt lõi (Visual Flow):</b> 
            <br>             
            <code style="font-size: 0.8rem; background-color: white; padding: 5px; border: 1px dashed #f9a825;">[1. Subject (Chart Type)]</code>
            <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
            <code style="font-size: 0.8rem; background-color: white; padding: 5px; border: 1px dashed #f9a825;">[2. Finite Verb]</code>
            <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
            <code style="font-size: 0.8rem; background-color: white; padding: 5px; border: 1px dashed #f9a825;">[3. Object (Topic + namely...)]</code>
            <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
            <code style="font-size: 0.8rem; background-color: white; padding: 5px; border: 1px dashed #f9a825;">[4. Place]</code>
            <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
            <code style="font-size: 0.8rem; background-color: white; padding: 5px; border: 1px dashed #f9a825;">[5. Time]</code>
        </div>
    </li>

    <li>
        <div style="background-color:#f8f9fa; border:1px solid #bdc3c7; border-radius:5px; padding:15px; margin-bottom:15px;">
            <strong style="color:#d35400;">📚 KHO CÔNG CỤ XÂY DỰNG INTRO:</strong>
            
            <!-- KHỐI 1: CHỦ NGỮ & ĐỘNG TỪ -->
            <details style="margin-top: 15px;">
                <summary style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; cursor: pointer; font-weight: bold;">
                    📍 BƯỚC 1 & 2: SUBJECT & FINITE VERB
                </summary>
                <div style="padding: 10px; background-color: white; border: 1px solid #fdf2e9;">
                    <b>● Subject (Chart Type):</b> <i>The bar chart / The line graph / The pie charts...</i>
                    <br><b>● Finite Verb:</b> <i>illustrates / compares / gives information about...</i>
                    <br>⚠️ <b>Hòa hợp:</b> 1 biểu đồ dùng <b>-s</b>, nhiều biểu đồ <b>không -s</b>.
                </div>
            </details>

            <!-- KHỐI 2: OBJECT/TOPIC -->
            <details style="margin-top: 10px;">
                <summary style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; cursor: pointer; font-weight: bold;">
                    📍 BƯỚC 3: OBJECT (PHẦN TRĂM/SỐ LƯỢNG + TOPIC)
                </summary>
                <div style="padding: 10px; background-color: white; border: 1px solid #fdf2e9;">
                    <b>● Bắt đầu bằng:</b> <i>The proportion of / The percentage of / The number of...</i>
                    <br><b>● Topic Paraphrase:</b> 
                    <ul style="margin: 5px 0;">
                        <li>...how much money was allocated to...</li>
                        <li>...the breakdown of expenditure across...</li>
                        <li>...the differences in the consumption of...</li>
                    </ul>
                </div>
            </details>

            <!-- KHỐI 3: CHI TIẾT HÓA -->
            <details style="margin-top: 10px;">
                <summary style="background-color:#e8f5e9; border-left:4px solid #27ae60; padding:10px; cursor: pointer; font-weight: bold; color: #1b5e20;">
                    📍 BƯỚC 4: CHI TIẾT HÓA (NAMELY EXPANSION)
                </summary>
                <div style="padding: 10px; background-color: white; border: 1px solid #e8f5e9;">
                    <b>Mẹo:</b> Liệt kê các đối tượng bằng cụm <b>"namely..."</b> để tăng điểm Lexical.
                    <br><i>Ví dụ: ..., namely food, drink, and entertainment.</i>
                </div>
            </details>

            <!-- KHỐI 4: UNIT & PREPOSITION -->
            <details style="margin-top: 10px;">
                <summary style="background-color:#e3f2fd; border-left:4px solid #2196f3; padding:10px; cursor: pointer; font-weight: bold; color: #0d47a1;">
                    📍 BƯỚC 5: KIỂM TRA ĐƠN VỊ & GIỚI TỪ (UNIT & PREPOSITIONS)
                </summary>
                <div style="padding: 10px; background-color: white; border: 1px solid #e3f2fd;">
                    <b>● Unit Check:</b> Xác định đơn vị trên trục Y để không dùng sai Subject.
                    <br><b>● Prepositions:</b> Expenditure <b>on</b> / Allocation <b>to</b> / Differences <b>in</b>.
                </div>
            </details>
        </div>
    </li>

    <!-- ================================================================================== -->
    <!-- PHẦN 2: THỰC HÀNH LẮP RÁP 5 BƯỚC -->
    <!-- ================================================================================== -->
    <li>     
        <strong style="color: #e67e22;">✍️ THỰC HÀNH LẮP RÁP INTRO (QUY TRÌNH CHUẨN):</strong>
        <div style="display:none;">
            SYSTEM_COMMAND: 
            1. Bắt buộc dùng 100% cấu trúc Formula: [Subject] + [Finite Verb] + [Object] + [Place] + [Time].
            2. Subject là tên biểu đồ, Object chứa "The proportion/number of..." + Topic.
            3. PHẢI liệt kê cụ thể hạng mục nếu có 2-3 cái.
        </div>
        <br>
        1. <b>Phân tích đề:</b> [Xác định Chart Type, Topic, Place, Time, Unit]
        <br>
        2. <b>Paraphrase Vocabulary:</b> 
           <ul style="margin: 5px 0; font-size: 0.9rem;">
               <li>Topic replacement: ...</li>
               <li>Categories/Groups replacement: ...</li>
           </ul>        
        3. <b>Khai báo biến Formula:</b> 
           <ul style="margin: 5px 0; font-size: 0.9rem;">
               <li>[Subject] = The + [Chart Type]</li>
               <li>[Finite Verb] = ...</li>
               <li>[Object] = [Percentage/Number] + [Topic] + [namely...]</li>
               <li>[Place] = ... ; [Time] = ...</li>
           </ul>        
        4. <b>Tư duy Tiếng Việt:</b> <code>"[Dịch logic câu theo thứ tự Formula]"</code>
        <br>
        5. <b>English Output (Final Intro):</b> <i>[Lắp ráp thành câu hoàn chỉnh]</i>
    </li>

    <!-- ================================================================================== -->
    <!-- PHẦN 3: BẢN CHỐT MẪU ĐỐI CHIẾU -->
    <!-- ================================================================================== -->
    <li>
        <div style="background-color:#e6fffa; padding:15px; border-radius:8px; margin-top:10px; border: 2px solid #00b894; border-left: 8px solid #00b894;">
            <b>📝 Bản chốt mở bài chuẩn (Final Polished Intro):</b><br>
            <div style="margin-top:5px; font-style: italic; color: #004d40;">
            [SYSTEM_PROMPT: Tổng hợp các biến đã chọn thành 01 câu duy nhất. 
             YÊU CẦU: Tuân thủ tuyệt đối trật tự: [Subject] + [Finite Verb] + [Object] + [Place] + [Time]. 
             Dùng đúng từ vựng đã paraphrase ở Bước 2. Đây là bản đáp án chuẩn.]
            </div>
        </div>
    </li>
</ul>

                    ### 2. **"overview_guide" (Ranking + Gap / Static Comparison):**

<ul>
    <!-- PHẦN 1: LOGIC CỐT LÕI (VISUAL FLOW) -->
    <li>
        <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">             
            <b>► Logic cốt lõi (Visual Flow):</b> 
            <br><code style="font-size: 0.8rem; background-color: white; padding: 5px; border: 1px dashed #f9a825;">[Overall] ➔ [1. Gom Đặc điểm Lớn nhất/Nhỏ nhất] ➔ [2. Diễn giải ý nghĩa (Widely used) / Big Picture (Sự phụ thuộc)]</code>
        </div>
    </li>

    <!-- PHẦN 2: KHO MẪU CÂU (SỔ XUỐNG) -->
    <li>
        <div style="background-color:#f8f9fa; border:1px solid #bdc3c7; border-radius:5px; padding:15px; margin-bottom:15px;">
            <strong style="color:#c0392b;">⛔ 3 LƯU Ý "SỐNG CÒN" KHI VIẾT OVERVIEW:</strong>
            <br><i>1. Tuyệt đối <b>KHÔNG</b> đưa số liệu cụ thể (%, số lượng) vào phần này.</i>
            <br><i>2. <b>KHÔNG</b> nhắc đến nhóm hạng mục phụ gộp chung <b>"Other / Others"</b>.</i>
            <br><i>3. <b>TRÁNH</b> dùng từ mang ý kiến cá nhân/cảm xúc như <b>"popular, favorite"</b> (Nên dùng: commonly/widely used).</i>

            <details style="margin-top: 15px;">
                <summary style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; cursor: pointer; font-weight: bold; color: #d35400;">
                    📚 KHO MẪU CÂU OVERVIEW (STATIC) - CÓ ĐÁNH SỐ ID
                </summary>
                <div style="padding: 15px; background-color: white; border: 1px solid #fdf2e9; border-top: none;">
                    
                    <b style="color: #e67e22;">► KHỐI 1: Cấu trúc cơ bản (So sánh Lớn nhất / Nhỏ nhất):</b>
                    <br><i>(Tư duy: Tìm ra điểm lớn nhất của từng đối tượng và nối bằng "while/whereas")</i>
                    <br>📍 <b>ID 1.1 (So sánh trực tiếp):</b> <code>Overall, [Subject A] [V: produced/consumed/spent...] the most [Topic] from/on [Category A], while [Subject B]...</code>
                    <br>📍 <b>ID 1.2 (Lấy Data làm chủ ngữ):</b> <code>Overall, [Category A] accounted for / made up the greatest percentage of [Topic] in [Subject].</code>
                    
                    <br><br><b style="color: #e67e22;">► KHỐI 2: Nâng cấp - Diễn tả ý nghĩa con số (Meaning):</b>
                    <br><i>(Tư duy: Đổi từ việc "đọc số lớn nhất" sang "bản chất hành động" - phổ biến nhất, dùng nhiều nhất)</i>
                    <br>📍 <b>ID 2.1:</b> "Overall, <b>[Category A]</b> was the most widely/commonly used source/type of <b>[Topic]</b> in <b>[Subject]</b>."
                    <br>📍 <b>ID 2.2:</b> "<b>[Category A]</b> was the primary choice / primary source of <b>[Topic]</b> for <b>[Subject]</b>."
                    
                    <br><br><b style="color: #e67e22;">► KHỐI 3: Nâng cấp - Bức tranh tổng thể & Sự phụ thuộc (Big Picture / Gap):</b>
                    <br><i>(Tư duy: Gom nhóm các hạng mục nhỏ thành [Broader Category] để thấy sự chênh lệch lớn)</i>
                    <br>📍 <b>ID 3.1:</b> "<b>Furthermore, [Subject A]</b> relied more heavily on / showed a higher dependence on <b>[Broader Category]</b> than <b>[Subject B]</b>."
                    <br>📍 <b>ID 3.2:</b> "While <b>[Subject A]</b> showed a higher preference for <b>[Broader Category 1]</b>, <b>[Subject B]</b> relied predominantly on <b>[Broader Category 2]</b>."
                </div>
            </details>
        </div>
    </li>

    <!-- PHẦN 3: THỰC HÀNH LẮP RÁP 5 BƯỚC CHUẨN -->
    <li>     
        <strong style="color: #e67e22;">✍️ THỰC HÀNH LẮP RÁP OVERVIEW (DÂY CHUYỀN 5 BƯỚC):</strong>
        <div style="display:none;">
            SYSTEM_COMMAND: 
            1. Bắt buộc dùng 100% cấu trúc của[Mẫu gốc] đã chọn. Chỉ thay thế các biến trong [ ].
            2. Tuyệt đối KHÔNG đưa số liệu (%) vào bài viết. Không dùng "Other" hay "popular".
            3. Expert Version = Câu 1 (từ ID Khối 1 hoặc 2) + Câu 2 (từ ID Khối 3 - Big Picture).
        </div>
        <br>1. 🔴 <b>Phân tích đề:</b>[Xác định hạng mục lớn nhất & Gom nhóm Big Picture]
        <br>2. 📥 <b>Mẫu gốc (ID):</b>[AI in 2 mẫu gốc sẽ sử dụng, ví dụ ID 2.1 + ID 3.1]
        <br>3. 🧩 <b>Khai báo biến:</b> [Subject] = ... ; [Category A] = ... ; [Broader Category] = ...
        <br>4. 🧠 <b>Tư duy Tiếng Việt:</b> <code>"[Dịch ý tưởng Overview]"</code>
        <br>5. 📝 <b>English Output:</b> <i>[Lắp ráp thành đoạn văn 2 câu hoàn chỉnh]</i>
    </li>

    <!-- PHẦN 4: KẾT QUẢ ĐỐI CHIẾU -->
    <li>
        <div style="background-color:#fff3e0; padding:15px; border-radius:8px; margin-top:10px; border: 2px solid #ff9f43; border-left: 8px solid #ff9f43;">
            <b>📝 Bản Overview chuẩn (Expert Version):</b><br>
            <div style="margin-top:5px; font-family: 'Georgia', serif; font-style: italic; color: #5d4037;">[AI tổng hợp 2 câu ở Bước 5 thành đoạn văn. Yêu cầu: Đúng mẫu 100%, có sự kết nối mượt mà giữa ý "Lớn nhất" và "Big Picture".]
            </div>
        </div>
    </li>
</ul>

                    ### 3. **"body1_guide" (Thân bài 1 - Nhóm Lớn Nhất / Nổi Bật Nhất):**
<ul>
    <!-- ================================================================================== -->
    <!-- MA TRẬN ĐA DẠNG HÓA CẤU TRÚC -->
    <!-- ================================================================================== -->
    <li>
        <div style="background-color:#e8f5e9; padding:15px; border:1px solid #4caf50; border-radius:5px; margin-bottom:15px;">
            <strong style="color:#2e7d32;">🎨 KỸ THUẬT ĐA DẠNG HÓA CHỦ NGỮ (SUBJECT VARIATION):</strong>
            <br><i>(Examiner yêu cầu: Không được viết 2 câu liên tiếp có cùng kiểu chủ ngữ. Hãy xoay tua 4 loại sau:)</i>
            <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 10px; margin-top: 10px; font-size: 0.9rem;">
                <div style="background-color:white; padding:8px; border-radius:4px; border:1px dashed #a5d6a7;">
                    <b>1. Category làm chủ ngữ:</b><br><i>"<b>Gas</b> accounted for 30%..."</i>
                </div>
                <div style="background-color:white; padding:8px; border-radius:4px; border:1px dashed #a5d6a7;">
                    <b>2. Từ chỉ Lượng làm chủ ngữ:</b><br><i>"<b>The proportion of Gas</b> was 30%..."</i>
                </div>
                <div style="background-color:white; padding:8px; border-radius:4px; border:1px dashed #a5d6a7;">
                    <b>3. Xếp hạng làm chủ ngữ:</b><br><i>"<b>The leading source</b> was Gas..."</i>
                </div>
                <div style="background-color:white; padding:8px; border-radius:4px; border:1px dashed #a5d6a7;">
                    <b>4. Nơi chốn/Người làm chủ ngữ:</b><br><i>"<b>France</b> produced 30%..."</i>
                </div>
            </div>
        </div>
    </li>

    <!-- ================================================================================== -->
    <!-- CÔNG THỨC TỔNG QUÁT -->
    <!-- ================================================================================== -->
    <li><b>Logic chọn nhóm (Grouping Strategy):</b> 
        <br>AI thực hiện: Chọn các hạng mục (Categories) có số liệu **LỚN NHẤT** hoặc chiếm **TỶ TRỌNG ÁP ĐẢO**.
        <br>🚩 <b>QUY TẮC CHỐNG LẶP DỮ LIỆU (STATIC):</b> 
        <br>1. Mỗi con số (Data point) chỉ được xuất hiện DUY NHẤT một lần trong cả bài.
        <br>2. Nếu đã dùng số liệu đó để so sánh ở câu trước, câu sau chỉ mô tả đặc tính hoặc gap (khoảng cách).
    </li>
    <li>
        <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-bottom:15px;">             
            <b>► Cấu trúc chuẩn (Formula):</b> 
            <br><i>(Quy trình 3 giai đoạn chuẩn hóa cho Body 1 Static)</i>
            <br>             
            <code style="font-size: 0.8rem; background-color: white; padding: 5px; border: 1px dashed #f9a825;">[1. Top Ranking]</code>
            <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
            <code style="font-size: 0.8rem; background-color: white; padding: 5px; border: 1px dashed #f9a825;">[2. Math Comparison]</code>
            <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
            <code style="font-size: 0.8rem; background-color: white; padding: 5px; border: 1px dashed #f9a825;">[3. Complex Integration]</code>
        </div>
    </li>

    <!-- ================================================================================== -->
    <!-- KHO MẪU CÂU SỔ XUỐNG -->
    <!-- ================================================================================== -->
    <li>
        <div style="background-color:#f8f9fa; border:1px solid #bdc3c7; border-radius:5px; padding:15px; margin-bottom:15px;">
            <strong style="color:#d35400;">📚 KHO MẪU CÂU BODY 1 (STATIC):</strong>
            <br><i>(Lưu ý: <b>[V_Tense]</b> = Chia động từ theo Năm của đề bài)</i>         
            
            <!-- KHỐI 1: SỔ XUỐNG -->
            <details style="margin-top: 15px;">
                <summary style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; cursor: pointer; font-weight: bold; color: #d35400;">
                    📍 KHỐI 1 - ĐIỂM XUẤT PHÁT & SO SÁNH ĐỐI CHIẾU (TOP RANKING)
                </summary>
                <div style="padding: 15px; background-color: white; border: 1px solid #fdf2e9; border-top: none;">
                    <div style="margin-bottom: 10px;">
                        <b>► Logic cốt lõi (Visual Flow):</b> 
                        <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[LINKER]</code>
                        <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
                        <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[CAT A + DATA A]</code>
                        <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
                        <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[COMPARISON WITH CAT B / CHART 2]</code>
                    </div>
                    <div style="font-size: 0.9rem; line-height: 1.6;">
                        <b>📍 [Mẫu 1.1 - Mệnh đề Which (Cross-Comparison)]:</b><br>
                        <code>Looking at the charts in more detail, <b>[Cat A]</b> accounted for <b>[Data A]</b> in <b>[Place 1]</b>, <b>which was [Gap Data] higher/lower than</b> the figure for <b>[Place 2]</b>.</code>
                        <br><br><b>📍 [Mẫu 1.2 - Cấu trúc While/Whereas (Đối lập)]:</b><br>
                        <code><b>While [Cat A]</b> was the dominant category in <b>[Place 1]</b> at <b>[Data A]</b>, the corresponding figure for <b>[Place 2]</b> was significantly lower at <b>[Data B]</b>.</code>
                        <br><br><b>📍 [Mẫu 1.3 - Cấu trúc Followed by (Thứ bậc)]:</b><br>
                        <code><b>[Cat A]</b> registered the highest proportion of <b>[Data A]</b>, <b>closely followed by [Cat B]</b>, which stood at <b>[Data B]</b>.</code>
                        <br><br><b>📍 [Mẫu 1.4 - Cấu trúc Compared to (Đối chiếu trực tiếp)]:</b><br>
                        <code>In <b>[Place 1]</b>, <b>[Cat A]</b> stood at <b>[Data A]</b>, <b>compared to a much [higher/lower] figure of [Data B]</b> for <b>[Cat B]</b>.</code>
                        <br><br><b>📍 [Mẫu 1.5 - Cấu trúc Similarity (Tương đồng)]:</b><br>
                        <code><b>[Cat A]</b> was the primary <b>[Topic]</b> in both regions, accounting for <b>[Data A]</b> and <b>[Data B]</b> respectively.</code>
                    </div>
                </div>
            </details>

            <!-- KHỐI 2: SỔ XUỐNG -->
            <details style="margin-top: 10px;">
                <summary style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; cursor: pointer; font-weight: bold; color: #d35400;">
                    📍 KHỐI 2 - DIỄN BIẾN THỨ CẤP & SO SÁNH TỶ LỆ
                </summary>
                <div style="padding: 15px; background-color: white; border: 1px solid #fdf2e9; border-top: none;">
                    <div style="margin-bottom: 10px;">
                        <b>► Logic cốt lõi (Visual Flow):</b> 
                        <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[LINKER]</code>
                        <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
                        <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[VARYING SUBJECT]</code>
                        <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
                        <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[FEATURE + DATA A]</code>
                        <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
                        <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[COMPARE WITH B]</code>
                    </div>
                    <div style="font-size: 0.9rem; line-height: 1.6;">
                        <b>📍 [Mẫu 2.1 - Cấu trúc Gấp lần (Math Focus)]:</b><br>
                        <code><b>[Linker]</b>, the second most popular category was <b>[Cat A]</b> at <b>[Data A]</b>, a figure nearly <b>[double / triple] that of [Cat B]</b>.</code>
                        <br><br><b>📍 [Mẫu 2.2 - Chủ ngữ hành động (Topic Focus)]:</b><br>
                        <code><b>[Linker]</b>, <b>[Cat A] consumption / spending</b> stood at <b>[Data A]</b>, which was <b>significantly [higher/lower] than</b> the proportion of <b>[Cat B]</b> at <b>[Data B]</b>.</code>
                        <br><br><b>📍 [Mẫu 2.3 - Cấu trúc Whereas (Đối lập song song)]:</b><br>
                        <code><b>[Linker]</b>, <b>[Data A]</b> of <b>[Total]</b> was attributed to <b>[Cat A]</b>, <b>whereas the figure for [Cat B] was much lower</b> at only <b>[Data B]</b>.</code>
                        <br><br><b>📍 [Mẫu 2.4 - Cấu trúc Appositive (Đồng cách)]:</b><br>
                        <code><b>[Linker]</b>, <b>[Cat A]</b>, the second leading source, accounted for <b>[Data A]</b>, <b>surpassing [Cat B] by a margin of [Gap Data]</b>.</code>
                        <br><br><b>📍 [Mẫu 2.5 - Cấu trúc Similarity (Bám đuổi)]:</b><br>
                        <code><b>[Linker]</b>, a comparable pattern was seen in <b>[Cat A]</b> with <b>[Data A]</b>, <b>only [Gap Data] [more/less] than</b> the figure recorded for <b>[Cat B]</b>.</code>
                    </div>
                </div>
            </details>

            <!-- KHỐI 3: SỔ XUỐNG -->
            <details style="margin-top: 10px;">
                <summary style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; cursor: pointer; font-weight: bold; color: #d35400;">
                    📍 KHỐI 3 - GOM NHÓM & TỔNG HỢP SỐ LIỆU (DATA PACKING)
                </summary>
                <div style="padding: 15px; background-color: white; border: 1px solid #fdf2e9; border-top: none;">
                    <div style="margin-bottom: 10px;">
                        <b>► Logic cốt lõi (Visual Flow):</b> 
                        <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[LINKER]</code>
                        <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
                        <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[GROUPED SUBJECTS]</code>
                        <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
                        <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[COLLECTIVE / RESPECTIVE DATA]</code>
                        <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
                        <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[FINAL COMPARISON]</code>
                    </div>
                    <div style="font-size: 0.9rem; line-height: 1.6;">
                        <b>📍 [Mẫu 3.1 - Cấu trúc Collectively (Gom tổng)]:</b><br>
                        <code><b>[Linker]</b>, <b>[Cat C]</b> and <b>[Cat D]</b> collectively accounted for <b>[Total Data]</b>, which was still lower than the figure for <b>[Cat A]</b>.</code>
                        <br><br><b>📍 [Mẫu 3.2 - Cấu trúc Respective (Liệt kê đối xứng)]:</b><br>
                        <code><b>[Linker]</b>, the proportions of <b>[Cat C]</b> and <b>[Cat D]</b> stood at <b>[Data C]</b> and <b>[Data D]</b> respectively, showing a marked difference compared to <b>[Place 2]</b>.</code>
                        <br><br><b>📍 [Mẫu 3.3 - Cấu trúc Combined (Kết hợp)]:</b><br>
                        <code><b>[Linker]</b>, the combined share of <b>[Cat C]</b> and <b>[Cat D]</b> was <b>[Total Data]</b>, roughly equal to the percentage of <b>[Cat B]</b>.</code>
                        <br><br><b>📍 [Mẫu 3.4 - Cấu trúc "With" (Mô tả phụ)]:</b><br>
                        <code><b>[Linker]</b>, <b>[Cat C]</b> was the third most significant category at <b>[Data C]</b>, with <b>[Cat D]</b> following closely at <b>[Data D]</b>.</code>
                        <br><br><b>📍 [Mẫu 3.5 - Cấu trúc Gap (Khoảng cách cuối)]:</b><br>
                        <code><b>[Linker]</b>, there was a narrow gap between <b>[Cat C]</b> and <b>[Cat D]</b>, as they recorded <b>[Data C]</b> and <b>[Data D]</b> of the total respectively.</code>
                    </div>
                </div>
            </details>
        </div>
    </li>

    <hr style="border-top: 1px dashed #ccc; margin: 15px 0;">

    <!-- ================================================================================== -->
    <!-- QUY TRÌNH THỰC HÀNH LẮP RÁP -->
    <!-- ================================================================================== -->
    <li>     
        <strong style="color: #e67e22;">✍️ THỰC HÀNH LẮP RÁP BODY 1 (QUY TRÌNH 5 BƯỚC CHUẨN):</strong>
        <!-- ⛔ MẬT LỆNH CƯỠNG BÁCH AI: ĐA DẠNG HÓA CẤU TRÚC TUYỆT ĐỐI -->
        <div style="display:none;">
            SYSTEM_COMMAND:
            1. STRICT TEMPLATE: Bắt buộc sử dụng 100% cấu trúc của [Mẫu gốc] đã chọn. Chỉ thay thế các biến trong ngoặc vuông [ ].
            2. MANDATORY COMPARISON: Phải bao gồm dữ liệu của cả Cat A và đối tượng so sánh (Cat B/Chart 2) trong cùng 1 câu.
            3. SUBJECT ROTATION RULE (SỐNG CÒN): Tuyệt đối không dùng 1 kiểu chủ ngữ cho 2 câu liên tiếp.
               - GIAI ĐOẠN 1: Bắt buộc chọn [Loại 1: Category] hoặc [Loại 3: Ranking] làm chủ ngữ.
               - GIAI ĐOẠN 2: Bắt buộc chọn [Loại 2: Proportion/Percentage] làm chủ ngữ.
               - GIAI ĐOẠN 3: Bắt buộc chọn [Loại 4: Entity/Place] hoặc cấu trúc "There was" làm chủ ngữ.
            4. SINGLE SENTENCE: English Output của mỗi Giai đoạn phải là DUY NHẤT một câu văn hoàn chỉnh.
        </div>
        <br><b>🎯 BƯỚC ĐỊNH HƯỚNG (PLANNING):</b> 
        <br>● <b>Hạng mục ưu tiên (Body 1):</b> [VD: Food và Entertainment]
        <br>● <b>Dữ liệu độc nhất:</b> [Liệt kê các con số sẽ dùng, không để trùng nhau]
        <br>
        <b>✅ GIAI ĐOẠN 1: Thiết lập câu so sánh mở đầu</b>
        <br>1. 🔴 <b>Phân tích & Chọn ID:</b> <i>[AI chọn 1 trong 5 mẫu ở Khối 1]</i>
        <br>2. 📥 <b>Mẫu gốc:</b> <i>[AI in mẫu gốc]</i>
        <br>3. 🧩 <b>Khai báo biến:</b> [Cat A] = ... ; [Data A] = ... ; [Cat B/Place 2] = ... ; [Data B] = ...
        <br>4. 🧠 <b>Tư duy Tiếng Việt:</b> <code>"[AI viết câu tư duy có tính chất so sánh]"</code>
        <br>5. 📝 <b>English Output:</b> <i>[AI xuất 01 câu duy nhất đúng chuẩn mẫu đã chọn]</i>
        
        <br><br><b>✅ GIAI ĐOẠN 2: Mô tả đặc tính thứ cấp và đối chiếu</b>
        <br>1. 🔴 <b>Phân tích & Chọn ID:</b> <i>[AI chọn mẫu ở Khối 2]</i>
        <br>2. 📥 <b>Mẫu gốc:</b> <i>[AI in mẫu gốc]</i>
        <br>3. 🧩 <b>Khai báo biến:</b> [Linker] = ... ; [Cat A/B] = ... ; [Data A/B] = ... ; [Feature] = ...
        <br>4. 🧠 <b>Tư duy Tiếng Việt:</b> <code>"[AI viết câu tư duy: Thằng tiếp theo là A chiếm X, gấp đôi thằng B...]"</code>
        <br>5. 📝 <b>English Output:</b> <i>[AI xuất 01 câu duy nhất]</i>

        <br><br><b>✅ GIAI ĐOẠN 3: Gom nhóm các hạng mục còn lại và chốt dữ liệu Body 1</b>
        <br>1. 🔴 <b>Phân tích & Chọn ID:</b> <i>[AI chọn mẫu ở Khối 3]</i>
        <br>2. 📥 <b>Mẫu gốc:</b> <i>[AI in mẫu gốc]</i>
        <br>3. 🧩 <b>Khai báo biến:</b> [Linker] = ... ; [Cat C/D] = ... ; [Data C/D] = ... ; [Comparison] = ...
        <br>4. 🧠 <b>Tư duy Tiếng Việt:</b> <code>"[AI viết câu tư duy: Cuối cùng, C và D gộp lại chiếm X%, vẫn thấp hơn thằng A...]"</code>
        <br>5. 📝 <b>English Output:</b> <i>[AI xuất 01 câu duy nhất]</i>
    </li>

    <!-- ================================================================================== -->
    <!-- KẾT QUẢ CUỐI CÙNG -->
    <!-- ================================================================================== -->
    <li>
        <div style="background-color:#fff8e1; padding:20px; border-radius:8px; margin-top:10px; border: 2px solid #ffa502; border-left: 8px solid #ffa502;">
            <strong style="color:#d35400; font-size: 1.1rem;">📝 KẾT QUẢ BODY 1 (THE POLISHED STATIC PARAGRAPH):</strong>
            <br><i style="font-size: 0.85rem; color: #7f8c8d;">(Đây là đoạn văn hoàn chỉnh tập trung vào xếp hạng và so sánh tỷ trọng điểm nổi bật)</i>
            
            <div style="margin-top:10px; font-family: 'Georgia', serif; line-height: 1.6; color: #2c3e50; background-color: white; padding: 15px; border-radius: 5px; border: 1px inset #eee;">
            [AI trình bày đoạn văn Body 1 hoàn chỉnh tại đây. Lưu ý: In đậm các cụm từ so sánh xếp hạng và từ vựng nâng Band.]
            </div>

            <!-- EXAMINER'S QUALITY AUDIT -->
            <div style="margin-top:10px; font-size: 0.8rem; color: #27ae60;">
                <b>✅ Examiner's Quality Audit (Static Comparison):</b>
                <ul style="margin: 0; padding-left: 20px;">
                    <li>Ưu tiên mô tả hạng mục cao nhất ngay câu đầu: <b>Xác nhận</b></li>
                    <li>Sử dụng cấu trúc so sánh hơn/gấp lần giữa các đối tượng: <b>Có</b></li>
                    <li>Xoay vòng chủ ngữ (Category vs Proportion): <b>Đạt chuẩn</b></li>
                    <li>Sử dụng giới từ Static (stood at/a figure of) chính xác: <b>Có</b></li>
                </ul>
            </div>
        </div>
    </li>
</ul>

                    ### 4. **"body2_guide" (Thân bài 2 - Các hạng mục còn lại / Nhóm nhỏ):**
<ul>
    <!-- ================================================================================== -->
    <!-- CHIẾN THUẬT GOM NHÓM (GROUPING STRATEGY) -->
    <!-- ================================================================================== -->
    <li>
        <div style="background-color:#e1f5fe; padding:15px; border:1px solid #0288d1; border-radius:5px; margin-bottom:15px;">
            <strong style="color:#01579b;">🧩 CHIẾN THUẬT GOM NHÓM (GROUPING STRATEGY):</strong>
            <br><i>(Để tránh lỗi liệt kê, hãy áp dụng quy tắc 3 chiều sau:)</i>
            <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 10px; margin-top: 10px; font-size: 0.9rem;">
                <div style="background-color:white; padding:8px; border-radius:4px; border:1px dashed #0288d1;">
                    <b>1. Gom theo tính tương đồng:</b><br>Nếu 2-3 mục có số liệu xấp xỉ nhau, hãy dùng cấu trúc <i>"comparable levels"</i> hoặc <i>"range from... to..."</i>.
                </div>
                <div style="background-color:white; padding:8px; border-radius:4px; border:1px dashed #0288d1;">
                    <b>2. Gom theo tổng lượng:</b><br>Dùng <i>"collectively"</i> hoặc <i>"combined share"</i> để cộng dồn các mục quá nhỏ thành một con số ấn tượng.
                </div>
            </div>
        </div>
    </li>

    <!-- ================================================================================== -->
    <!-- CÔNG THỨC TỔNG QUÁT -->
    <!-- ================================================================================== -->
    <li>
        <div style="background-color:#eaf2f8; border-left:4px solid #2980b9; padding:10px; margin-bottom:15px;">             
            <b>► Cấu trúc chuẩn (Formula):</b> 
            <br><i>(Quy trình wrap-up cho các hạng mục còn lại)</i>
            <br>             
            <code style="font-size: 0.8rem; background-color: white; padding: 5px; border: 1px dashed #2980b9;">[1. Transition]</code>
            <span style="color:#2980b9; font-weight:bold;"> ➔ </span>
            <code style="font-size: 0.8rem; background-color: white; padding: 5px; border: 1px dashed #2980b9;">[2. Grouping/Similarities]</code>
            <span style="color:#2980b9; font-weight:bold;"> ➔ </span>
            <code style="font-size: 0.8rem; background-color: white; padding: 5px; border: 1px dashed #2980b9;">[3. The Smallest Category]</code>
        </div>
    </li>

    <!-- ================================================================================== -->
    <!-- KHO MẪU CÂU SỔ XUỐNG -->
    <!-- ================================================================================== -->
    <li>
        <div style="background-color:#f8f9fa; border:1px solid #bdc3c7; border-radius:5px; padding:15px; margin-bottom:15px;">
            <strong style="color:#2980b9;">📚 KHO MẪU CÂU BODY 2 (STATIC):</strong>
            <br><i>(Lưu ý: Dùng màu xanh dương để phân biệt với nhóm dẫn đầu Body 1)</i>         
            
            <!-- KHỐI 1: CHUYỂN ĐOẠN -->
            <details style="margin-top: 15px;">
                <summary style="background-color:#eaf2f8; border-left:4px solid #2980b9; padding:10px; cursor: pointer; font-weight: bold; color: #2980b9;">
                    📍 KHỐI 1 - CÂU CHUYỂN ĐOẠN & ĐỐI CHIẾU (TRANSITION)
                </summary>
                <div style="padding: 15px; background-color: white; border: 1px solid #eaf2f8; border-top: none;">
                    <div style="margin-bottom: 10px; font-size: 0.9rem;">
                        <b>► Logic cốt lõi:</b> <code>[TRANSITION LINKER] ➔ [CAT C & D IDENTITY] ➔ [CONTRAST WITH BODY 1]</code>
                    </div>
                    <div style="font-size: 0.9rem; line-height: 1.6;">
                        <b>📍 [Mẫu 4.1 - Chuyển hướng]:</b><br>
                        <code><b>Turning to the remaining categories</b>, the figures for <b>[Cat C]</b> and <b>[Cat D]</b> were significantly lower, at <b>[Data C]</b> and <b>[Data D]</b> respectively.</code>
                        <br><br><b>📍 [Mẫu 4.2 - Đối lập hoàn toàn]:</b><br>
                        <code><b>In stark contrast to the aforementioned sectors</b>, <b>[Cat C]</b> represented a much smaller share of the total, standing at only <b>[Data C]</b>.</code>
                        <br><br><b>📍 [Mẫu 4.3 - Nhóm trung bình]:</b><br>
                        <code><b>As for the mid-range categories</b>, <b>[Cat C]</b> and <b>[Cat D]</b> recorded <b>comparable levels</b> of <b>[Data C]</b> and <b>[Data D]</b>.</code>
                    </div>
                </div>
            </details>

            <!-- KHỐI 2: GOM NHÓM -->
            <details style="margin-top: 10px;">
                <summary style="background-color:#eaf2f8; border-left:4px solid #2980b9; padding:10px; cursor: pointer; font-weight: bold; color: #2980b9;">
                    📍 KHỐI 2 - GOM NHÓM & KHOẢNG DỮ LIỆU (GROUPING)
                </summary>
                <div style="padding: 15px; background-color: white; border: 1px solid #eaf2f8; border-top: none;">
                    <div style="margin-bottom: 10px; font-size: 0.9rem;">
                        <b>► Logic cốt lõi:</b> <code>[COMBINED SUBJECTS] ➔ [COLLECTIVE VERB] ➔ [DATA RANGE]</code>
                    </div>
                    <div style="font-size: 0.9rem; line-height: 1.6;">
                        <b>📍 [Mẫu 5.1 - Collectively]:</b><br>
                        <code><b>[Cat C]</b> and <b>[Cat D]</b> <b>collectively accounted for</b> a mere <b>[Total Data]</b> of the overall <b>[Topic]</b>.</code>
                        <br><br><b>📍 [Mẫu 5.2 - Combined share]:</b><br>
                        <code>The <b>combined share</b> of the remaining categories, including <b>[Cat C]</b> and <b>[Cat D]</b>, stood at <b>[Total Data]</b>.</code>
                        <br><br><b>📍 [Mẫu 5.3 - Range]:</b><br>
                        <code>The proportions of <b>[Cat C, D, and E]</b> were relatively minor, <b>ranging from [Lowest Data] to [Highest Data]</b>.</code>
                    </div>
                </div>
            </details>

            <!-- KHỐI 3: THIỂU SỐ -->
            <details style="margin-top: 10px;">
                <summary style="background-color:#eaf2f8; border-left:4px solid #2980b9; padding:10px; cursor: pointer; font-weight: bold; color: #2980b9;">
                    📍 KHỐI 3 - HẠNG MỤC THẤP NHẤT (THE MINORITY)
                </summary>
                <div style="padding: 15px; background-color: white; border: 1px solid #eaf2f8; border-top: none;">
                    <div style="margin-bottom: 10px; font-size: 0.9rem;">
                        <b>► Logic cốt lõi:</b> <code>[THE LEAST POPULAR] ➔ [NEGLIGIBLE/MINORITY] ➔ [DATA]</code>
                    </div>
                    <div style="font-size: 0.9rem; line-height: 1.6;">
                        <b>📍 [Mẫu 6.1 - Negligible]:</b><br>
                        <code>Finally, the figure for <b>[Cat E]</b> was <b>negligible</b>, representing only <b>[Data E]</b> of the total.</code>
                        <br><br><b>📍 [Mẫu 6.2 - Least significant]:</b><br>
                        <code><b>[Cat E]</b> was the <b>least significant category</b>, with a proportion of just <b>[Data E]</b>.</code>
                        <br><br><b>📍 [Mẫu 6.3 - Small minority]:</b><br>
                        <code>Only a <b>small minority</b> of <b>[People/Objects]</b> belonged to the <b>[Cat E]</b> group (<b>[Data E]</b>).</code>
                    </div>
                </div>
            </details>
        </div>
    </li>

    <hr style="border-top: 1px dashed #ccc; margin: 15px 0;">

    <!-- ================================================================================== -->
    <!-- QUY TRÌNH THỰC HÀNH LẮP RÁP BODY 2 -->
    <!-- ================================================================================== -->
    <li>     
        <strong style="color: #2980b9;">✍️ THỰC HÀNH LẮP RÁP BODY 2 (QUY TRÌNH 5 BƯỚC CHUẨN):</strong>
        <div style="display:none;">
            SYSTEM_COMMAND: 
            1. Bắt buộc dùng Transition Linker để nối với Body 1.
            2. Sử dụng kỹ thuật gom nhóm (collectively/range) để bài viết súc tích.
            3. Output duy nhất 01 câu cho mỗi giai đoạn.
        </div>
        <br>
        <br><b>✅ GIAI ĐOẠN 1: Câu chuyển đoạn và mô tả nhóm hạng mục trung bình</b>
        <br>1. 🔴 <b>Phân tích & Chọn ID:</b> <i>[AI chọn mẫu ở Khối 1 - Body 2]</i>
        <br>2. 📥 <b>Mẫu gốc:</b> <code>[AI in mẫu gốc]</code>
        <br>3. 🧩 <b>Khai báo biến:</b> [Linker] = ... ; [Cat C/D] = ... ; [Data C/D] = ...
        <br>4. 🧠 <b>Tư duy Tiếng Việt:</b> <code>"[AI viết câu chuyển đoạn: Chuyển sang các mục còn lại, C và D thấp hơn nhiều...]"</code>
        <br>5. 📝 <b>English Output:</b> <i>[AI xuất 01 câu duy nhất đúng chuẩn mẫu]</i>
        
        <br><br><b>✅ GIAI ĐOẠN 2: Gom nhóm hoặc so sánh tương đồng các hạng mục nhỏ</b>
        <br>1. 🔴 <b>Phân tích & Chọn ID:</b> <i>[AI chọn mẫu ở Khối 2 - Body 2]</i>
        <br>2. 📥 <b>Mẫu gốc:</b> <code>[AI in mẫu gốc]</code>
        <br>3. 🧩 <b>Khai báo biến:</b> [Cat C/D/E] = ... ; [Total Data/Range] = ...
        <br>4. 🧠 <b>Tư duy Tiếng Việt:</b> <code>"[AI viết câu tư duy: Hai hạng mục X và Y gộp lại chỉ chiếm một lượng nhỏ là...]"</code>
        <br>5. 📝 <b>English Output:</b> <i>[AI xuất 01 câu duy nhất đúng chuẩn mẫu]</i>

        <br><br><b>✅ GIAI ĐOẠN 3: Chốt hạng mục thấp nhất bài (The Minority)</b>
        <br>1. 🔴 <b>Phân tích & Chọn ID:</b> <i>[AI chọn mẫu ở Khối 3 - Body 2]</i>
        <br>2. 📥 <b>Mẫu gốc:</b> <code>[AI in mẫu gốc]</code>
        <br>3. 🧩 <b>Khai báo biến:</b> [Cat E] = ... ; [Data E] = ... ; [Feature] = (negligible/minority)
        <br>4. 🧠 <b>Tư duy Tiếng Việt:</b> <code>"[AI viết câu tư duy: Cuối cùng, hạng mục Z chiếm tỷ lệ không đáng kể, chỉ ở mức...]"</code>
        <br>5. 📝 <b>English Output:</b> <i>[AI xuất 01 câu duy nhất đúng chuẩn mẫu]</i>
    </li>

    <!-- ================================================================================== -->
    <!-- KẾT QUẢ CUỐI CÙNG BODY 2 -->
    <!-- ================================================================================== -->
    <li>
        <div style="background-color:#e1f5fe; padding:20px; border-radius:8px; margin-top:10px; border: 2px solid #03a9f4; border-left: 8px solid #03a9f4;">
            <strong style="color:#01579b; font-size: 1.1rem;">📝 KẾT QUẢ BODY 2 (THE POLISHED STATIC PARAGRAPH):</strong>
            <br><i style="font-size: 0.85rem; color: #7f8c8d;">(Đoạn văn hoàn chỉnh cho các nhóm còn lại, nhấn mạnh tính tổng hợp)</i>
            
            <div style="margin-top:10px; font-family: 'Georgia', serif; line-height: 1.6; color: #2c3e50; background-color: white; padding: 15px; border-radius: 5px; border: 1px inset #eee;">
            [AI trình bày đoạn văn Body 2 hoàn chỉnh tại đây. Lưu ý: In đậm các cụm từ gom nhóm và từ nối chuyển đoạn.]
            </div>

            <!-- EXAMINER'S QUALITY AUDIT -->
            <div style="margin-top:10px; font-size: 0.8rem; color: #0288d1;">
                <b>✅ Examiner's Quality Audit (Grouping & Contrast):</b>
                <ul style="margin: 0; padding-left: 20px;">
                    <li>Sử dụng Transition Linker để dẫn dắt: <b>Xác nhận</b></li>
                    <li>Áp dụng kỹ thuật gom nhóm (Collectively/Combined): <b>Có</b></li>
                    <li>Sử dụng từ vựng miêu tả mức độ thấp (Negligible/Minority): <b>Đạt chuẩn</b></li>
                    <li>Bao phủ toàn bộ các hạng mục còn lại của đề bài: <b>Có</b></li>
                </ul>
            </div>
        </div>
    </li>
</ul>
                    # =================================================================
                    # 🔵 TRƯỜNG HỢP 2: DẠNG "CHANGE OVER TIME" (Line, Bar, Table, Pie nhiểu năm)
                    # (Tư duy cốt lõi: Trend (Xu hướng) & Speed (Tốc độ thay đổi))
                    # =================================================================

                    1. **"intro_guide" (Paraphrase):**                    
    <ul>
        <li><b>Cấu trúc chuẩn (Formula):</b> 
        <br><code>[Subject] + [Finite Verb] + [Object/Topic] + [Place] + [Time]</code>.</li>        
    
        <li><b>Subject (Lưu ý quan trọng):</b>
            <br>- <b>Xác định đúng chủ thể:</b> <i>[Xác định chính xác cái gì thay đổi]</i>.
            <br>- <b>Đơn vị trong bài này là:</b> <i>[Điền đơn vị cụ thể của bài, VD: million dollars / tonnes / %]</i>.
            <br>- <b>Tuyệt đối không đưa đơn vị tính vào chủ ngữ.</b> Ví dụ: Không viết <i>"The chart shows [Đơn vị của bài]..."</i> mà phải viết <i>"The chart shows the amount/number/proportion of..."</i>.
            <br>- <b>Hòa hợp chủ ngữ - động từ:</b> Nếu 1 biểu đồ dùng <i>shows/illustrates</i>. Nếu nhiều biểu đồ dùng <i>show/illustrate</i>.</li>
    
        <li><b>Cách đổi Chủ ngữ & Topic (The "What") cho bài này:</b>
            <br>- <b>Từ vựng gốc trong đề:</b> "<i>[Trích cụm từ gốc trong đề bài]</i>"
            <br>- <b>Gợi ý Paraphrase 1:</b> <i>[Viết phương án paraphrase 1. VD: The amount of money spent on...]</i>
            <br>- <b>Gợi ý Paraphrase 2:</b> <i>[Viết phương án paraphrase 2. VD: How much money was allocated to...]</i>
            <br><i>(Lưu ý: chọn từ Spending/Number/Percentage phù hợp).</i></li>
    
        <li><b>Verb (Động từ khuyên dùng):</b>
            <br><i>illustrates / gives information about / compares the data on / presents information about</i>.</li>
    
        <li><b>Time Paraphrase (Thời gian: [Năm đầu] - [Năm cuối]):</b>
            <br>- Cách 1: <i>Between [Năm đầu] and [Năm cuối]</i>.
            <br>- Cách 2: <i>Over a period of [Số năm] years starting from / commencing in [Năm đầu]</i>.</li>    
                
        <li><div style="background-color:#e6fffa; padding:10px; border-radius:5px; margin-top:5px; border-left: 4px solid #00b894;">
            <b>📝 Nội dung mẫu (Sample Intro):</b><br>                             
            <i>[Viết câu Introduction theo hướng dẫn đã phân tích]</i>
        </div></li>
    </ul>

                    2. **"overview_guide" (Trend + Ranking):**
                        - <ul>
                         <!-- ================================================================================== -->
                         <!-- PHẦN 1: KHO MẪU CÂU (CHẾ ĐỘ IN ẤN - KHÔNG TƯ DUY) -->                         
                         <!-- ================================================================================== -->
                         <li>
                             <div style="background-color:#f8f9fa; border:1px solid #bdc3c7; border-radius:5px; padding:15px; margin-bottom:15px;">
                             <strong style="color:#d35400;">📚 KHO MẪU CÂU:</strong>
                                                         
                             <!-- KHỐI 1: TỔNG QUÁT -->
                             <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">
                                 <b>► 1. Cấu trúc chuẩn (Formula):</b>
                                 <br><code>Overall, &#91;Sentence 1: Trends&#93;. In addition, &#91;Sentence 2: Highlights&#93;.</code>
                             </div>
                                                        
                             <div style="background-color:#ffebee; border:1px solid #ef5350; padding:10px; margin-bottom:10px; border-radius:4px;">
                                 <strong style="color:#c62828;">⛔ LƯU Ý VỀ ĐỘNG TỪ (VERB RESTRICTIONS):</strong>
                                 <br>1. Với chủ ngữ là <b>SỐ LIỆU</b> (The figure / The percentage):
                                 <br>• <b>Nên dùng:</b> <i>increased, decreased, rose, fell</i> (Trực tiếp).
                                 <br>• <b>Chấp nhận:</b> <i>experienced, underwent</i> (Trải qua).
                                 <br>• <b>HẠN CHẾ:</b> <i>saw, witnessed</i> (Vì con số không có mắt để nhìn).
                                 
                                 <br>2. Với chủ ngữ là <b>HẠNG MỤC / NƠI CHỐN</b> (The UK / Farming):
                                 <br>👉 <b>Phải dùng:</b> <i>saw / witnessed / experienced ... <b>IN</b> ...</i>
                             </div>
                             
                             <!-- KHỐI 2: TRENDS (ĐÃ KIỂM TRA KỸ LƯỠNG - BUG FREE) -->
                             <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">
                                 <b>► 2. Các mẫu câu Xu hướng (Trends):</b>
                                 <br>✅ <b>Loại 1: Đồng loạt Tăng/Giảm (Same Direction)</b>
                                 <br><i>(Dùng khi tất cả cùng Tăng hoặc cùng Giảm)</i>                                 
                                 <br>• <b>Số liệu:</b> "It is clear that the <b>total [number/amount/percentage] of &#91;Topic&#93;</b> increased/decreased <b>over the period shown</b>."
                                 <br>• <b>Hạng mục:</b> "It is noticeable that <b>&#91;Place/Category&#93;</b> witnessed an upward/downward trend <b>IN</b> <b>&#91;Topic&#93;</b> <b>throughout the given period</b>."                                 
                                 <br>✅ <b>Loại 2: Xu hướng ngược (Mix / Opposite)</b>
                                 <br><i>(CHỈ DÙNG khi A Tăng còn B Giảm - Ngược chiều hoàn toàn)</i>                                 
                                 <br>• <b>Số liệu:</b> "It is clear that while the <b>figures for &#91;Category A&#93;</b> increased, the opposite was true for <b>&#91;Category C&#93;</b> <b>over the given period</b>."
                                 <br>• <b>Hạng mục:</b> "It is noticeable that while <b>&#91;Category A&#93;</b> saw an upward trend <b>IN</b> <b>&#91;Topic&#93;</b>, <b>&#91;Category C&#93;</b> experienced a decline <b>during the period shown</b>."                                 
                                 <br>✅ <b>Loại 3: Ngoại lệ (Exception)</b>
                                 <br><i>(Dùng khi đa số Tăng, chỉ có 1 cái Giảm/Ổn định)</i>                                 
                                 <br>• <b>Mẫu cơ bản:</b> "The figures for most categories increased, <b>with the exception of &#91;Category C&#93;</b>, <b>over the given period</b>."
                                 <br>• <b>Mẫu nâng cao:</b> "<b>With the exception of &#91;Category C&#93;</b>, all other categories <b>witnessed</b> an upward/downward trend <b>over the period shown</b>."                                 
                                 <br>✅ <b>Loại 4: Ổn định/Dao động (Stability/Fluctuation)</b>
                                 <br><i>(Dùng khi có đường đi ngang hoặc dao động mạnh)</i>
                                 <br>"It is clear that while <b>&#91;Category A&#93;</b> changed significantly, the figure for <b>&#91;Category B&#93;</b> remained relatively stable <b>over the given period</b>."
                             </div>

                             <!-- KHỐI 3: HIGHLIGHTS (LIỆT KÊ ĐỦ 3 LOẠI) -->
                             <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">
                                 <b>► 3. Các mẫu câu Điểm nổi bật ( Highlights):</b>
                                 <br><i>(Phân tích đề bài và hình ảnh để chọn Highlights loại nào?, và chỉ hiển thị công thức)</i>                               
                                 <br>✅ <b>Loại 1: Cao nhất/Thấp nhất (Ranking)</b>
                                 <br><i>(Dùng khi có 1 đường luôn nằm trên hoặc nằm dưới các đường khác)</i>
                                 <br>"<b>&#91;Category A&#93;</b> consistently had the highest figures throughout the period."                                 
                                 <br>✅ <b>Loại 2: Biến động lớn nhất (Biggest Change)</b>
                                 <br><i>(Dùng khi có 1 đường tăng/giảm mạnh nhất so với bọn còn lại)</i>
                                 <br>"<b>&#91;Category B&#93;</b> witnessed the most dramatic change."                               
                                 <br>✅ <b>Loại 3: Soán ngôi (Ranking Shift)</b>
                                 <br><i>(Dùng khi các đường cắt nhau)</i>
                                 <br>"The <b>figure for &#91;Category A&#93;</b> overtook <b>that of &#91;Category B&#93;</b> to become the dominant category."                               
                             </div>
                         </li>
                         
                         <hr style="border-top: 1px dashed #ccc; margin: 15px 0;">

                         <!-- ================================================================================== -->
                         <!-- PHẦN 2: CHIẾN THUẬT CHỌN MẪU (STRICT MATCHING & DRAFTING) -->                         
                         <!-- ================================================================================== -->
                         <li>
                             <div style="background-color:#e1f5fe; padding:15px; border-radius:8px; border-left: 5px solid #0288d1; margin-bottom: 15px;">
                                 <strong style="color:#01579b; font-size: 1.1rem;">🔍 PHÂN TÍCH BÀI NÀY (Selection & Drafting):</strong>
                                 <br><i>(Học sinh nhìn hình để chọn Mẫu cho phù hợp)</i>
                                 
                                 <!-- ⛔ MẬT LỆNH: CẤM AI SÁNG TẠO (Người dùng không thấy đoạn này) -->
                                 <div style="display:none;">
                                     SYSTEM_COMMAND:
                                     1.  **SCAN TREND:** Phân tích kỹ xu hướng thực tế của hình ảnh.
                                     2.  **MATCH OR BUILD:** 
                                        - Nếu có ID khớp > 80%: Hãy COPY ID đó và điền từ.
                                        - Nếu KHÔNG có ID nào khớp: BẮT BUỘC sử dụng tư duy "LEGO" - Ghép 2 mẫu lại với nhau (như cách bạn ghép 2.3 và 2.4).
                                     3.  **HYBRID RULE:** Khi ghép mẫu, phải đảm bảo tính nhất quán của chủ ngữ. 
                                        - Sai: "There was a rise... and it surpassed..." (Lỗi mix giữa There was và It).
                                        - Đúng: "The figure experienced a rise, surpassing... and ending at..." (Sử dụng mệnh đề quan hệ/rút gọn để ghép).
                                     4.  **EXCEPTION:** Nếu trend là "Đi ngang rồi mới tăng" (mà mẫu không có), hãy dùng cấu trúc: "[Subject] remained stable before witnessing a rise to [Data]."
                                 </div>

                                 <br><br><b>1. Với Xu hướng (Trend - Câu 1):</b>
                                 <ul style="margin-top:5px; list-style-type: none; padding-left: 10px;">
                                     <li>👀 <b>Quan sát (Observation):</b> <i>[AI MÔ TẢ NGẮN GỌN DẤU HIỆU. Ví dụ: "Hai đường đi ngược chiều nhau."]</i></li>
                                     <li>🧩 <b>Khớp với Mẫu số:</b> <b>[AI ĐIỀN LOẠI]</b>.</li>
                                     
                                     <!-- Phần viết câu: Dùng thẻ code để nổi bật kết quả -->
                                     <br>👉 <b>Viết câu Trends:</b> 
                                     <br><code style="color:#d35400; background-color:#fff; padding: 2px 5px; border-radius: 4px;">[AI điền dữ liệu vào mẫu. YÊU CẦU: Copy y nguyên cấu trúc mẫu, KHÔNG ĐƯỢC sửa từ vựng]</code>
                                 </ul>

                                 <br><b>2. Với Điểm nổi bật (Highlight - Câu 2):</b>
                                  <ul style="margin-top:5px; list-style-type: none; padding-left: 10px;">
                                     <li>👀 <b>Quan sát (Observation):</b> <i>[AI MÔ TẢ NGẮN GỌN DẤU HIỆU]</i></li>
                                     <li>🧩 <b>Khớp với Mẫu số:</b> <b>[AI ĐIỀN LOẠI]</b>.</li>
                                     
                                     <!-- Phần viết câu -->
                                     <br>👉 <b>Viết câu Highlights:</b> 
                                     <br><code style="color:#d35400; background-color:#fff; padding: 2px 5px; border-radius: 4px;">[AI điền dữ liệu vào mẫu. YÊU CẦU: Copy y nguyên cấu trúc mẫu, KHÔNG ĐƯỢC sửa từ vựng]</code>
                                 </ul>
                             </div>                            
                         </li>

                         <!-- PHẦN 3: TỪ VỰNG -->
                         <li><b>🔑 TỪ VỰNG GỢI Ý (Vocabulary):</b>
                             <br>- <b>Verbs:</b> <i>witnessed a downward trend / saw a significant rise</i>.
                             <br>- <b>Adjectives:</b> <i>volatile (biến động) / stable (ổn định)</i>.
                             <br>- <b>Paraphrase bài này:</b> <i>[AI liệt kê 3 từ vựng sát với chủ đề bài viết]</i>
                         </li>

                         <!-- PHẦN 4: BÀI MẪU -->
                         <li><div style="background-color:#fff3e0; padding:15px; border-radius:8px; margin-top:10px; border-left: 5px solid #ff9f43;">
                             <b>📝 Nội dung mẫu (Sample Overview):</b><br>
                             <div style="margin-top:5px; font-style: italic; color: #5d4037;">
                             [AI hãy viết đoạn Overview hoàn chỉnh. <b>QUAN TRỌNG:</b> Chỉ được sử dụng đúng cấu trúc câu đã chọn ở PHẦN 2]
                             </div>
                         </div></li>
                       </ul>

### 🕒 PHẦN BỔ TRỢ: KHO THỜI GIAN (ĐẶT TRƯỚC BODY 1)

   - <ul>
     <!-- ================================================================================== -->
     <!-- GLOBAL MODULE: KHO THỜI GIAN VẠN NĂNG -->
     <!-- SYSTEM_ALERT: AI tham chiếu bảng này. QUY TẮC: Không dùng "In [Year]" 2 lần liên tiếp. -->
     <!-- ================================================================================== -->
     <li>
         <div style="background-color:#e3f2fd; border:1px solid #2196f3; border-radius:5px; padding:15px; margin-bottom:15px;">
             <strong style="color:#0d47a1;">⏰ KHO THỜI GIAN VẠN NĂNG (UNIVERSAL TIME MENU):</strong>
             <br><i>(Hướng dẫn: Chọn 1 cụm từ dưới đây để điền vào biến số <b>[Time]</b>. Hãy xoay tua giữa Cột 1 và Cột 2)</i>
             
             <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 10px; margin-top: 10px; font-size: 0.9rem;">
                 <!-- CỘT 1: THỜI GIAN TUYỆT ĐỐI (Dùng cho câu đầu đoạn) -->
                 <div style="background-color:white; padding:8px; border-radius:4px; border:1px dashed #2196f3;">
                     <b style="color:#1565c0;">1. Mốc Tuyệt đối (Absolute):</b>
                     <ul style="margin:5px 0; padding-left:15px;">
                         <li>In <b>[Year]</b> (VD: In 1999)</li>
                         <li>At the start of the period</li>
                         <li>In the final year / By <b>[Year]</b></li>
                         <li>Between <b>[Year]</b> and <b>[Year]</b></li>
                     </ul>
                 </div>

                 <!-- CỘT 2: THỜI GIAN TƯƠNG ĐỐI (Dùng để nối câu - QUAN TRỌNG) -->
                 <div style="background-color:white; padding:8px; border-radius:4px; border:1px dashed #2196f3;">
                     <b style="color:#1565c0;">2. Tương đối & Trình tự (Relative):</b>
                     <br><i style="font-size: 0.8rem; color:#d84315;">(Dùng cái này để tránh lặp lại năm)</i>
                     <ul style="margin:5px 0; padding-left:15px;">
                         <li><b>[X] years later</b> (VD: 10 years later)</li>
                         <li><b>Over the next [X] years</b></li>
                         <li><b>Following this / Thereafter</b> (Sau đó)</li>
                         <li>In the subsequent years</li>
                     </ul>
                 </div>
             </div>
         </div>
     </li>
   </ul>
### 📏 PHẦN BỔ TRỢ 2: KHO GIỚI TỪ & ƯỚC LƯỢNG (PRECISION & PREPOSITIONS)
*Học sinh thường mất điểm GRA (Ngữ pháp) cực kỳ nặng vì sai giới từ.*
     <li>
         <div style="background-color:#fff3e0; border:1px solid #ff9800; border-radius:5px; padding:15px; margin-bottom:15px;">
             <strong style="color:#e65100;">🎯 KHO GIỚI TỪ & ƯỚC LƯỢNG (THE PRECISION KIT):</strong>
             <br><i>(Quy tắc: Sai giới từ = Sai bản chất số liệu. Hãy dùng bảng này để kiểm tra biến <b>[Data]</b>)</i>
             
             <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 10px; margin-top: 10px; font-size: 0.9rem;">
                 <!-- CỘT 1: GIỚI TỪ THẦN THÁNH -->
                 <div style="background-color:white; padding:8px; border-radius:4px; border:1px dashed #ff9800;">
                     <b style="color:#bf360c;">1. Giới từ (Prepositions):</b>
                     <ul style="margin:5px 0; padding-left:15px;">
                         <li><b>Stood at / Started at:</b> Đứng tại (Mốc đầu)</li>
                         <li><b>Rose TO [X]:</b> Tăng ĐẾN mức X (Đích đến)</li>
                         <li><b>Rose BY [X]:</b> Tăng THÊM một khoảng X</li>
                         <li><b>A peak OF [X]:</b> Đạt đỉnh ở mức X</li>
                         <li><b>Stabilized AT [X]:</b> Ổn định tại mức X</li>
                     </ul>
                 </div>

                 <!-- CỘT 2: TỪ ƯỚC LƯỢNG -->
                 <div style="background-color:white; padding:8px; border-radius:4px; border:1px dashed #ff9800;">
                     <b style="color:#bf360c;">2. Ước lượng (Approximation):</b>
                     <br><i style="font-size: 0.8rem; color:#d84315;">(Dùng khi số liệu không nằm trên vạch kẻ)</i>
                     <ul style="margin:5px 0; padding-left:15px;">
                         <li><b>Just under / Slightly below:</b> Dưới ... một chút</li>
                         <li><b>Just over / Slightly above:</b> Trên ... một chút</li>
                         <li><b>Approximately / Roughly:</b> Khoảng chừng</li>
                         <li><b>Nearly / Almost:</b> Gần như (Chưa chạm tới)</li>
                     </ul>
                 </div>
             </div>
         </div>
     </li>

### 🔄 PHẦN BỔ TRỢ 3: KHO CHỦ NGỮ XOAY VÒNG (SUBJECT ROTATION MENU)
*Lỗi lặp từ "The figure for..." ở mọi câu là nguyên nhân khiến Band Lexical Resource không thể lên 7.0.*
     <li>
         <div style="background-color:#f1f8e9; border:1px solid #4caf50; border-radius:5px; padding:15px; margin-bottom:15px;">
             <strong style="color:#1b5e20;">🔄 KHO CHỦ NGỮ XOAY VÒNG (SUBJECT ROTATION - VERSION 3.0):</strong>
             <br><i>(CHIẾN THUẬT: Đa dạng hóa thực thể để phá bỏ lời nguyền lặp từ "The figure for")</i>
             
             <div style="display: grid; grid-template-columns: 1fr 1fr 1fr; gap: 10px; margin-top: 10px; font-size: 0.85rem;">
                 
                 <!-- NHÓM 1: CON SỐ CHUNG (GENERIC NUMERICAL) -->
                 <div style="background-color:white; padding:8px; border-radius:4px; border:1px dashed #4caf50;">
                     <b style="color:#2e7d32;">1. Con số chung:</b>
                     <br><i style="font-size: 0.75rem;">(Dùng tối đa 2 lần/bài)</i>
                     <ul style="margin:5px 0; padding-left:15px;">
                         <li>The figure for <b>[Cat A]</b></li>
                         <li>The data for <b>[Cat A]</b></li>
                         <li>The percentage/proportion of...</li>
                         <li>The share of... (Rất hay)</li>
                     </ul>
                 </div>

                 <!-- NHÓM 2: CHỦ ĐỀ CHUYÊN BIỆT (TOPIC-SPECIFIC) - CỰC QUAN TRỌNG -->
                 <div style="background-color:#fffde7; padding:8px; border-radius:4px; border:1px dashed #fbc02d;">
                     <b style="color:#f57f17;">2. Theo chủ đề (High Band):</b>
                     <br><i style="font-size: 0.75rem;">(Biến danh từ thành chủ ngữ)</i>
                     <ul style="margin:5px 0; padding-left:15px;">
                         <li><b>[Activity]</b> participation rates</li>
                         <li><b>[Activity]</b> popularity / levels</li>
                         <li><b>[Product]</b> consumption</li>
                         <li><b>[Sector]</b> expenditure / spending</li>
                     </ul>
                 </div>

                 <!-- NHÓM 3: ĐỊNH DANH (IDENTITY - OBJECT) -->
                 <div style="background-color:#e1f5fe; padding:8px; border-radius:4px; border:1px dashed #03a9f4;">
                     <b style="color:#0288d1;">3. Định danh (Object):</b>
                     <br><i style="font-size: 0.75rem;">(Cấm đi kèm số liệu trực tiếp)</i>
                     <ul style="margin:5px 0; padding-left:15px;">
                         <li><b>[Cat A]</b> was the dominant...</li>
                         <li><b>[Cat A]</b> was the most popular...</li>
                         <li><b>[Cat A]</b> surpassed <b>[Cat B]</b>...</li>
                         <li>Working in the <b>[Cat A]</b> sector...</li>
                     </ul>
                 </div>
             </div>

             <!-- LỆNH CƯỠNG BÁCH CHO AI -->
             <div style="background-color:#ffebee; padding:10px; border-radius:4px; margin-top:10px; border-left: 4px solid #c62828;">
                <p style="font-size: 0.85rem; color: #c62828; margin: 0;">
                    ⚠️ <b>ANTI-REPETITION COMMAND:</b> 
                    <br>1. <b>CẤM</b> dùng "The figure for" hai lần liên tiếp.
                    <br>2. <b>PHẢI</b> luân phiên sử dụng ít nhất 01 chủ ngữ từ <b>Nhóm 2 (Topic-Specific)</b> trong mỗi đoạn thân bài để đạt Band 7.5+ Lexical Resource.
                    <br>3. <b>LỖI LOGIC:</b> "The figure for A was the dominant sector" (SAI). Hãy viết "A was the dominant sector" (ĐÚNG).
                </p>
             </div>
         </div>
     </li>
                    3. **"body1_guide" (Thân bài 1 - Nhóm Nổi bật / Biến động mạnh):**
   - <ul>
     <li><b>Logic chọn nhóm (Grouping Strategy):</b> AI thực hiện: Chọn 1 hoặc 2 hạng mục (Category) có <b>số liệu cao nhất</b> hoặc <b>xu hướng thay đổi mạnh nhất/tăng trưởng rõ rệt</b> để phân tích trước.</li>
     
     <li>
         <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">             
             <b>► Cấu trúc chuẩn (Formula):</b> 
             <br><i>(Quy trình 3 bước chuẩn hóa cho Body 1)</i>
             <br>             
             <code style="font-size: 0.8rem; background-color: white; padding: 5px; border: 1px dashed #f9a825;">[1. So sánh Start]</code>
             <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
             <code style="font-size: 0.8rem; background-color: white; padding: 5px; border: 1px dashed #f9a825;">[2. Trend & End (Cat A)]</code>
             <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
             <code style="font-size: 0.8rem; background-color: white; padding: 5px; border: 1px dashed #f9a825;">[3. So sánh & End (Cat B)]</code>
         </div>
     </li>

     <li>
         <div style="background-color:#f8f9fa; border:1px solid #bdc3c7; border-radius:5px; padding:15px; margin-bottom:15px;">
             <strong style="color:#d35400;">📚 KHO MẪU CÂU BODY 1:</strong>
             <br><i>(Lưu ý: Chọn từ Menu Thời Gian để điền vào <b>[Time]</b>)</i>         
             
             <!-- KHỐI 1: SỔ XUỐNG -->
             <details style="margin-top: 15px;">
                 <summary style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; cursor: pointer; font-weight: bold;">
                     📍 KHỐI 1 - ĐIỂM XUẤT PHÁT & SO SÁNH BAN ĐẦU
                 </summary>
                 <div style="padding: 10px; border: 1px solid #fdf2e9; border-top: none; background-color: white;">
                     <div style="margin-bottom: 10px;">
                         <b>► Logic cốt lõi (Visual Flow):</b> 
                         <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[TIME]</code>
                         <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
                         <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[DATA A]</code>
                         <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
                         <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[COMPARE B]</code>
                     </div>
                     <div style="margin-top:10px; font-size: 0.9rem;">
                         <b>📍 [Mẫu 1.1 - Cấu trúc While] (Dùng khi A và B khác biệt rõ):</b>                 
                         <br><code><b>At the start of the period</b>, <b>[Cat A]</b> was the dominant sector at <b>[Data A]</b>, <b>while the figure for [Cat B] was significantly lower at [Data B]</b>.</code> 
                         <br><br><b>📍 [Mẫu 1.2 - Mệnh đề quan hệ] (Dùng để nhấn mạnh chênh lệch):</b>                 
                         <br><code><b>In [Year]</b>, <b>[Cat A]</b> was the most popular category, but it was <b>subsequently overtaken by [Cat B]</b>, which surged to <b>[Data B]</b>.</code>
                         <br><br><b>📍 [Mẫu 1.3 - Rút gọn] (Dùng khi A lớn hơn B một chút):</b>
                         <br><code><b>[Time]</b>, <b>[Cat A]</b> was the dominant category at <b>[Data A]</b>, <b>closely followed by [Cat B] with [Data B]</b>.</code>                 
                         <br><br><b>📍 [Mẫu 1.4 - Giới từ Compared to] (Dùng để đối chiếu):</b>
                         <br><code><b>At the beginning of the period</b>, <b>[Cat A]</b> registered a figure of <b>[Data A]</b>, <b>compared to [Data B] for [Cat B]</b>.</code>                 
                         <br><br><b>📍 [Mẫu 1.5 - Tương đồng] (Dùng khi A = B hoặc xấp xỉ):</b>
                         <br><code><b>In the first year</b>, the figure for <b>[Cat A]</b> stood at <b>[Data A]</b>, <b>which was [identical/similar] to that of [Cat B]</b>.</code>                 
                         <br><br><b>📍 [Mẫu 1.6 - 3 CAT] (Dùng cho 3+ nhóm sát nhau):</b>
                         <br><code><b>At the beginning of the period</b>, <b>[Cat A], [Cat B] and [Cat C]</b> were clustered at significant levels, <b>ranging from [Lowest Data] to [Highest Data]</b>.</code>
                         <br><br><b>📍 [Mẫu 1.7 - Xuất phát 0] (Dùng cho dữ liệu bằng 0):</b>
                         <br><code><b>In [Year]</b>, <b>[Cat A]</b> was non-existent (stood at 0), <b>whereas [Cat B] was already established at [Data B]</b>.</code>
                         <br><br><b>📍 [Mẫu 1.8 - Xấp xỉ] (Dùng khi số liệu khó đọc):</b>
                         <br><code><b>At the start of the period</b>, <b>[Cat A]</b> stood at approximately <b>[Data A]</b>, <b>marginally [higher/lower] than the figure for [Cat B]</b>.</code>
                     </div>
                 </div>
             </details>

             <!-- KHỐI 2: SỔ XUỐNG -->
             <details style="margin-top: 10px;">
                 <summary style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; cursor: pointer; font-weight: bold;">
                     📍 KHỐI 2 - MIÊU TẢ TREND & ĐIỂM GÃY
                 </summary>
                 <div style="padding: 10px; border: 1px solid #fdf2e9; border-top: none; background-color: white;">
                     <div style="margin-bottom: 10px;">
                         <b>► Logic cốt lõi (Visual Flow):</b> 
                         <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[LINKER]</code>
                         <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
                         <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[TREND & MID-POINT]</code>
                         <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
                         <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[END DATA & TIME]</code>
                     </div>
                     <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">         
                         <b>► KHỐI 2: MIÊU TẢ TREND (DIỄN BIẾN):</b>
                         <br><i>(Lưu ý: <b>[Linker]</b> là các từ nối thời gian: Then, Subsequently, Following this, Over the next X years...)</i>             
                         <div style="margin-top:10px;">
                             <b>📍 [Mẫu 2.1 - Before V-ing] (Tăng rồi Giảm/Ngược lại):</b>
                             <br><code><b>[Linker]</b>, the figure increased to <b>[Peak Data]</b> in <b>[Year]</b>, <b>before falling to finish at</b> <b>[End Data]</b> <b>[End Time]</b>.</code>
                         </div>
                         <div style="margin-top:10px;">
                             <b>📍 [Mẫu 2.2 - Peak/Low & End] (Đạt Đỉnh/Đáy rồi đổi chiều):</b>
                             <br>• <b>Đỉnh:</b> <code><b>[Linker]</b>, it surged to <b>reach a peak of [Peak Data]</b> in <b>[Year]</b>, <b>and then dropped to [End Data]</b> <b>[End Time]</b>.</code>
                             <br>• <b>Đáy:</b> <code><b>[Linker]</b>, it plunged to <b>hit a low of [Low Data]</b> in <b>[Year]</b>, <b>before recovering to [End Data]</b> <b>[End Time]</b>.</code>
                         </div>
                         <div style="margin-top:10px;">
                             <b>📍 [Mẫu 2.3 - Soán ngôi & End] (Tăng và Vượt mặt):</b>
                             <br><code><b>[Linker]</b>, it rose significantly <b>to [Mid Data]</b>, <b>surpassing [Cat B]</b> to become the dominant category and <b>ending at [End Data]</b> <b>[End Time]</b>.</code>
                         </div>
                         <div style="margin-top:10px;">
                             <b>📍 [Mẫu 2.4 - Tăng/Giảm thẳng] (Xu hướng đơn giản):</b>
                             <br><code><b>[Linker]</b>, there was a sharp rise to <b>[End Data]</b> <b>[End Time]</b>, making it the highest figure at the end of the period.</code>
                         </div>
                         <div style="margin-top:10px;">
                             <b>📍 [Mẫu 2.5 - Finishing] (Mệnh đề rút gọn):</b>
                             <br><code><b>[Linker]</b>, it followed a steady <b>[upward/downward]</b> trend, <b>finishing the period at [End Data]</b> <b>[End Time]</b>.</code>
                         </div>
                         <div style="margin-top:10px;">
                             <b>📍 [Mẫu 2.6 - Dao động] (Zíc-zắc & Kết thúc):</b>
                             <br><code><b>[Linker]</b>, it showed a volatile pattern, <b>fluctuating between [Data 1] and [Data 2]</b>, before ending at <b>[End Data]</b> <b>[End Time]</b>.</code>
                         </div>
                         <div style="margin-top:10px;">
                             <b>📍 [Mẫu 2.7 - Ổn định cao] (Luôn đứng nhất):</b>
                             <br><code><b>[Linker]</b>, the figure <b>remained relatively stable</b> at <b>around [Data]</b> throughout the period, maintaining its leading position until <b>[End Time]</b>.</code>
                         </div>
                         <div style="margin-top:10px;">
                             <b>📍 [Mẫu 2.8 - Plateau] (Tăng/Giảm rồi đi ngang):</b>
                             <br><code><b>[Linker]</b>, it climbed rapidly to reach <b>[Data]</b> in <b>[Year]</b>, <b>after which it leveled off/plateaued</b> for the remainder of the period.</code>
                         </div>
                         <div style="margin-top:10px;">
                             <b>📍 [Mẫu 2.9 - Recovery] (Hồi phục sau khi giảm):</b>
                             <br><code><b>[Linker]</b>, after an initial drop to <b>[Low Data]</b>, the figure <b>recovered</b>, rising back to <b>[End Data]</b> <b>[End Time]</b>.</code>
                         </div>
                         <div style="margin-top:10px;">
                             <b>📍 [Mẫu 2.10 - Advanced Synthesis] (Biến động kép):</b>
                             <br><code><b>[Linker]</b>, <b>[Cat A]</b> saw a <b>[Sharp/Steady]</b> <b>[Rise/Fall]</b>, <b>[Surpassing/Overtaking] [Cat B]</b> to finish the period as the leader at <b>[End Data]</b> <b>[End Time]</b>.</code>
                         </div>
                         <div style="margin-top:10px; border-top: 1px dashed #ccc; padding-top: 10px;">
                             <b>📍 [Mẫu 2.11 - Sleeping Giant] (Đi ngang rồi mới Tăng/Giảm):</b>
                             <br><i>(Dùng khi giai đoạn đầu không đổi, giai đoạn sau mới biến động)</i>
                             <br><code><b>[Linker]</b>, the figure <b>remained stable at [Start Data]</b> until <b>[Year]</b>, <b>before [rising/falling] sharply to end at [End Data]</b> <b>[End Time]</b>.</code>
                         </div>
                         <div style="margin-top:10px;">
                             <b>📍 [Mẫu 2.12 - Losing Position] (Giảm và bị vượt mặt):</b>
                             <br><i>(Dùng khi đường đang cao bị tụt xuống dưới đường khác)</i>
                             <br><code><b>[Linker]</b>, it experienced a downward trend, <b>falling below [Cat B]</b> in <b>[Year]</b> and <b>continuing to drop to [End Data]</b> <b>[End Time]</b>.</code>
                         </div>
                     </div>
                 </div>
             </details>

             <!-- KHỐI 3: SỔ XUỐNG -->
             <details style="margin-top: 10px;">
                 <summary style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; cursor: pointer; font-weight: bold;">
                     📍 KHỐI 3 - SO SÁNH VỚI CÁC CATEGORY CÒN LẠI
                 </summary>
                 <div style="padding: 10px; border: 1px solid #fdf2e9; border-top: none; background-color: white;">
                     <div style="margin-bottom: 10px;">
                         <b>► Logic cốt lõi (Visual Flow):</b> 
                         <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[LINKER]</code>
                         <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
                         <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[COMPARE & TREND B]</code>
                         <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
                         <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[CONTINUE/MID-POINT]</code>
                         <span style="color:#f57f17; font-weight:bold;"> ➔ </span>
                         <code style="font-size: 0.9rem; background-color: white; padding: 2px 5px; border: 1px dashed #7f8c8d; border-radius: 3px;">[END DATA & TIME]</code>
                     </div>
                     <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">
                         <b>► KHỐI 3: SO SÁNH & KẾT THÚC (COMPARISON & END):</b>
                         <br><i>(Dùng để đối chiếu Cat B với Cat A đã tả ở trên - Đảm bảo tính liên kết chặt chẽ)</i>
                         
                         <div style="margin-top:10px; border-bottom: 1px dashed #ccc; padding-bottom: 10px;">
                             <b style="color: #27ae60;">● Nhóm 1: Đồng thuận & Bám đuổi (Following A)</b>
                             <br>📍 <b>[Mẫu 3.1 - Mirroring] (Phản chiếu y hệt):</b>
                             <br><code><b>Similarly</b>, <b>[Cat B]</b> mirrored this trajectory, <b>[rising/falling] to [Mid Data]</b> before <b>continuing its trend to finish at [End Data]</b> <b>[End Time]</b>.</code>                 
                             <br>📍 <b>[Mẫu 3.2 - Lagging behind] (Bám đuổi nhưng chậm hơn):</b>
                             <br><code><b>A similar, albeit more gradual, [rise/fall]</b> was observed in <b>[Cat B]</b>, with the figure <b>creeping [up/down] to [End Data]</b> by <b>[End Time]</b>.</code>
                         </div>
                         <div style="margin-top:10px; border-bottom: 1px dashed #ccc; padding-bottom: 10px;">
                             <b style="color: #c0392b;">● Nhóm 2: Đối lập & Nghịch thế (Opposing A)</b>
                             <br>📍 <b>[Mẫu 3.3 - In stark contrast] (Trái ngược hoàn toàn):</b>
                             <br><code><b>In stark contrast to [Cat A]</b>, the figure for <b>[Cat B]</b> followed a <b>downward/upward</b> path, <b>dropping/surging to [Mid Data]</b> and <b>ending the period at [End Data]</b>.</code>
                             <br>📍 <b>[Mẫu 3.4 - Standing still] (A biến động, B đứng im):</b>
                             <br><code><b>While [Cat A] saw significant changes</b>, <b>[Cat B]</b> remained <b>relatively stagnant/flat</b>, <b>hovering around [Data]</b> throughout the period until <b>[End Time]</b>.</code>
                         </div>
                         <div style="margin-top:10px; border-bottom: 1px dashed #ccc; padding-bottom: 10px;">
                             <b style="color: #2980b9;">● Nhóm 3: Thu hẹp hoặc Nới rộng khoảng cách</b>
                             <br>📍 <b>[Mẫu 3.5 - Closing the gap] (B đuổi kịp A):</b>
                             <br><code><b>Subsequently</b>, <b>[Cat B]</b> experienced a much sharper <b>[increase/decrease]</b>, <b>effectively closing the gap with [Cat A]</b> and finishing at <b>[End Data]</b>, just <b>[Number] units</b> apart.</code>
                             <br>📍 <b>[Mẫu 3.6 - Divergence] (B ngày càng xa A):</b>
                             <br><code><b>Following this</b>, the trend for <b>[Cat B]</b> diverged from <b>[Cat A]</b>, <b>[plummeting/soaring] to [End Data]</b>, which was <b>significantly [lower/higher]</b> than its counterpart.</code>
                         </div>
                         <div style="margin-top:10px;">
                             <b style="color: #8e44ad;">● Nhóm 4: Toán học & Tỷ lệ (Quantitative Relationship)</b>
                             <br>📍 <b>[Mẫu 3.7 - Double/Triple] (Gấp đôi/Gấp ba):</b>
                             <br><code><b>Regarding [Cat B]</b>, it followed a comparable <b>[up/down]</b> trend to end at <b>[End Data]</b> <b>[End Time]</b>, a figure <b>[approximately/nearly/exactly] [double / triple] that of [Cat A]</b>.</code>
                             <br>📍 <b>[Mẫu 3.8 - Fraction] (Một phần nhỏ):</b>
                             <br><code><b>Meanwhile</b>, the figure for <b>[Cat B]</b> <b>[increased/decreased]</b> marginally, <b>finishing at only [End Data]</b>, representing <b>merely a fraction of [Cat A]'s total</b>.</code>
                         </div>
                     </div>
                 </div>
             </details>
         </div>
     </li>
     
     <hr style="border-top: 1px dashed #ccc; margin: 15px 0;">

     <!-- ================================================================================== -->
     <!-- PHẦN 2: GHÉP CÂU HOÀN CHỈNH -->
     <!-- SYSTEM_ALERT: QUY TRÌNH "TEMPLATE RECALL" (GỌI LẠI MẪU).
          AI BẮT BUỘC PHẢI IN RA MẪU GỐC TRƯỚC KHI ĐIỀN TỪ. -->
     <!-- ================================================================================== -->
     <li>     
         <b>✍️ THỰC HÀNH LẮP RÁP BODY 1 (CHẾ ĐỘ COPY-PASTE):</b>
         <!-- ⛔ MẬT LỆNH CƯỠNG BÁCH AI (Dùng display:none để ẩn với người dùng nhưng AI vẫn đọc được) -->
         <div style="display:none;">
             SYSTEM_COMMAND: 
             1. Bắt buộc sử dụng 100% cấu trúc của [Mẫu gốc]. 
             2. Chỉ được phép thay thế các cụm trong ngoặc vuông [ ] bằng dữ liệu đã khai báo. 
             3. TUYỆT ĐỐI CẤM thêm thông tin, thêm mốc thời gian hoặc thêm câu văn nằm ngoài các biến số của mẫu.
             4. English Output phải là DUY NHẤT một câu văn hoàn chỉnh theo đúng mẫu gốc.
         </div>
         <br><i>(Quy trình: Chọn ID &rarr; In Mẫu Gốc &rarr; Khai báo &rarr; Điền vào chỗ trống)</i>
         <br>
         <br><b>✅ BƯỚC 1: Xử lý Câu mở đầu (Starting Point)</b>
         <br><i>(Mục tiêu: Thiết lập bối cảnh thời gian và vị thế ban đầu của các đối tượng)</i>         
         <br>- <b>🔴 Phân tích & Chọn ID:</b> <i>[AI quan sát tương quan số liệu (A>B, A=B, Sát nút, Tụ họp, hay Bằng 0) để chọn Mẫu từ 1.1 đến 1.8]</i>         
         <br>- <b>📥 Mẫu gốc:</b> <code>[AI copy y nguyên cấu trúc mẫu đã chọn]</code>         
         <br>- <b>🧩 Khai báo biến (Variable Declaration):</b>
           <br>&nbsp;&nbsp;+ <b>[Time]</b> = <i>(Chọn 1 cụm từ Menu Thời Gian)</i>
           <br>&nbsp;&nbsp;+ <b>[Cat A] & [Data A]</b> = <i>...</i>
           <br>&nbsp;&nbsp;+ <b>[Cat B] & [Data B]</b> = <i>...</i>
           <br>&nbsp;&nbsp;+ <b>[Comparison Word]</b> = <i>...</i>         
         <br>- <b>🧠 Tư duy Tiếng Việt (Teacher's Analysis):</b> 
         <br><i>(Giáo viên phân tích ý tưởng trước khi viết)</i>
         <br><code>"[AI viết câu tiếng Việt tương ứng. VD: Vào năm 1995, số lượng X đứng ở mức 100, trong khi con số này của Y là 50.]"</code>         
         <br>- <b>📝 Điền từ (English Output):</b> 
         <br><i>[AI thay thế các biến số vào mẫu gốc để tạo câu tiếng Anh hoàn chỉnh]</i>
         <br>
         <br><b>✅ BƯỚC 2: Xử lý Diễn biến Chủ thể A (Main Trend & End)</b>
         <br><i>(Mục tiêu: Mô tả hành trình từ điểm bắt đầu đến điểm kết thúc của hạng mục quan trọng nhất)</i>         
         <br>- <b>🔴 Phân tích & Chọn ID:</b> <i>[AI quan sát đường đi của Cat A (Tăng, giảm, đạt đỉnh, hay dao động...) để chọn Mẫu từ 2.1 đến 2.12]</i>         
         <br>- <b>📥 Mẫu gốc:</b> <code>[AI copy y nguyên cấu trúc mẫu đã chọn]</code>         
         <br>- <b>🧩 Khai báo biến (Variable Declaration):</b>
           <br>&nbsp;&nbsp;+ <b>[Linker]</b> = <i>(Chọn 1 cụm nối thời gian tương đối. VD: Subsequently / Over the next 10 years)</i>
           <br>&nbsp;&nbsp;+ <b>[Mid-point Data/Year]</b> = <i>(Số liệu tại điểm gãy hoặc năm đạt Đỉnh/Đáy)</i>
           <br>&nbsp;&nbsp;+ <b>[End Data]</b> = <i>(Số liệu tại năm cuối cùng)</i>
           <br>&nbsp;&nbsp;+ <b>[End Time]</b> = <i>(Mốc thời gian cuối cùng. VD: in the final year / by 2010)</i>         
         <br>- <b>🧠 Tư duy Tiếng Việt (Teacher's Analysis):</b> 
         <br><i>(Giáo viên phân tích diễn biến của Cat A)</i>
         <br><code>"[AI viết câu tiếng Việt tương ứng. VD: Sau đó, con số của X tăng mạnh và đạt đỉnh 200 vào năm 2000, trước khi giảm nhẹ xuống còn 180 vào cuối kỳ.]"</code>         
         <br>- <b>📝 Điền từ (English Output):</b> 
         <br><i>[AI lắp ghép các biến số vào mẫu để tạo câu mô tả diễn biến hoàn chỉnh cho Cat A]</i>         
         <br>
         <br><b>✅ BƯỚC 3: Xử lý So sánh & Diễn biến Chủ thể B (Comparison & End)</b>
         <br><i>(Mục tiêu: Đối chiếu hành trình của Cat B với Cat A và chốt số liệu cuối cùng)</i>         
         <br>- <b>🔴 Phân tích & Chọn ID:</b> <i>[AI nhìn tương quan Cat B so với Cat A (Ngược chiều, bám đuổi, hay gấp bao nhiêu lần...) để chọn Mẫu từ 3.1 đến 3.8]</i>         
         <br>- <b>📥 Mẫu gốc:</b> <code>[AI copy y nguyên cấu trúc mẫu đã chọn]</code>         
         <br>- <b>🧩 Khai báo biến (Variable Declaration):</b>
           <br>&nbsp;&nbsp;+ <b>[Linker/Comparison Word]</b> = <i>(VD: In contrast / Similarly / Mirroring this)</i>
           <br>&nbsp;&nbsp;+ <b>[Cat B]</b> = <i>(Tên hạng mục thứ hai trong nhóm)</i>
           <br>&nbsp;&nbsp;+ <b>[End Data]</b> = <i>(Số liệu năm cuối của Cat B)</i>
           <br>&nbsp;&nbsp;+ <b>[Relative Info]</b> = <i>(Thông tin so sánh thêm. VD: double that of A / closing the gap)</i>         
         <br>- <b>🧠 Tư duy Tiếng Việt (Teacher's Analysis):</b> 
         <br><i>(Giáo viên phân tích sự tương quan giữa B và A)</i>
         <br><code>"[AI viết câu tiếng Việt tương ứng. VD: Ngược lại, hạng mục Y đi theo xu hướng trái ngược khi giảm đều về mức 10 vào năm cuối, chỉ bằng một phần nhỏ so với X.]"</code>         
         <br>- <b>📝 Điền từ (English Output):</b> 
         <br><i>[AI hoàn thiện câu so sánh cuối cùng cho Body 1]</i>
     </li>
   
### 💎 PHẦN 3: BỘ CÔNG CỤ NÂNG BAND (THE EXAMINER'S ELITE TOOLKIT)
*Chỉ sử dụng khi đã nắm chắc các cấu trúc cơ bản. Nhắm thẳng vào tiêu chí Lexical Resource & Grammatical Range.*

<ul style="list-style-type: none; padding-left: 0;">
     <li>     
    <details style="cursor: pointer; border: 1px solid #bdc3c7; border-radius: 5px; padding: 10px; background-color: #f8f9fa;">
             <summary>
                 <strong style="color:#d35400; font-size: 1.1rem;">🚀 VŨ KHÍ NÂNG BAND (TOP-TIER ENHANCEMENTS):</strong>
             </summary>
             <br><i>(Các bạn có thể sử dụng ít nhất 1 cấu trúc nâng cao này cho mỗi đoạn thân bài)</i>
             <br>
             <ul style="margin-top:5px; line-height: 1.6;">
                 <!-- NHÓM 1: SỰ BIẾN THỂ DANH TỪ -->
                 <li style="margin-bottom: 10px;"><b>1. Biến thể Danh từ (Adjective + Noun):</b>
                     <br><i>(Thay vì viết "increased significantly", hãy dùng cấu trúc: There was a + adj + noun)</i>
                     <br>- <b>Mẫu:</b> <i>witnessed a <b>dramatic hike</b> / recorded a <b>slight dip</b> / experienced a <b>period of volatility</b>.</i>
                     <br>- <b>Advantage:</b> Giúp thay đổi hoàn toàn cấu trúc câu, ghi điểm Grammatical Range tuyệt đối.
                 </li>
                 <!-- NHÓM 2: CÁC ĐỘNG TỪ SANG CHẢNH -->
                 <li style="margin-bottom: 10px;"><b>2. Động từ "Sang chảnh":</b>
                     <br><i>(Dùng để thay thế cho các từ tăng/giảm đã quá nhàm chán)</i>
                     <br>- <b>Tăng vọt:</b> <i>rocket / soar / surge.</i>
                     <br>- <b>Chạm đáy:</b> <i>bottom out at / hit a rock bottom of.</i>
                     <br>- <b>Chiếm tỷ trọng:</b> <i>account for / constitute / comprise.</i>
                     <br>- <b>Hồi phục:</b> <i>bounce back / regain its position.</i>
                 </li>
                 <!-- NHÓM 3: TỪ VỰNG THEO CHỦ ĐỀ -->
                 <li style="margin-bottom: 10px;"><b>3. Topic Vocabulary (Nâng cao):</b>
                     <br>- <b>Chủ đề Chi tiêu (Money):</b> <i>expenditure, outlay, financial allocation, budget distribution.</i>
                     <br>- <b>Chủ đề Năng lượng/Hàng hóa:</b> <i>consumption, usage, throughput, production levels.</i>
                     <br>- <b>Chủ đề Dân số/Người:</b> <i>demographics, workforce, participants, residents.</i>
                 </li>
                 <!-- NHÓM 4: KỸ THUẬT CHÈN SỐ LIỆU PHỨC HỢP -->
                 <li><b>4. Chèn số liệu "nén" (Data Packing):</b>
                     <br><i>(Giúp đưa được nhiều thông tin vào một câu mà không bị rối)</i>
                     <br>- <b>Cấu trúc "With":</b> <i>...rose to 50%, <b>with</b> the most significant growth <b>occurring</b> in the final year.</i>
                     <br>- <b>Cấu trúc "Followed by":</b> <i>A was the highest at 50, <b>followed closely by</b> B and C at 45 and 40 respectively.</i>
                     <br>- <b>Cấu trúc "Relative":</b> <i>...ending at 100, <b>a figure which surpassed</b> all other categories.</i>
                 </li>
             </ul>
         </details>
     </li>
     </ul>

### PHẦN 4 - KẾT QUẢ BODY 1 (FINAL OUTPUT)
     <!-- ================================================================================== -->
     <!-- PHẦN 4: KẾT QUẢ BODY 1 (FINAL OUTPUT) -->
     <!-- SYSTEM_ALERT: 
          1. AI thực hiện nối 3 câu đã viết ở trên.
          2. Thực hiện "Cohesion Check": Đảm bảo các từ nối (Linkers) không bị lặp.
          3. Thực hiện "Vocabulary Upgrade": Sử dụng ít nhất 02 từ từ "Bộ nâng Band" vào đoạn văn.
     <!-- ================================================================================== -->
     <li>
         <div style="background-color:#fff8e1; padding:20px; border-radius:8px; margin-top:10px; border: 2px solid #ffa502; border-left: 8px solid #ffa502;">
             <strong style="color:#d35400; font-size: 1.1rem;">📝 KẾT QUẢ BODY 1 (THE POLISHED PARAGRAPH):</strong>
             <br><i style="font-size: 0.85rem; color: #7f8c8d;">(Đây là đoạn văn hoàn chỉnh đã được tinh chỉnh về mặt liên kết và từ vựng nâng cao)</i>
             
             <div style="margin-top:10px; font-family: 'Georgia', serif; line-height: 1.6; color: #2c3e50; background-color: white; padding: 15px; border-radius: 5px; border: 1px inset #eee;">
             [AI trình bày đoạn văn Body 1 hoàn chỉnh tại đây. 
              Lưu ý: In đậm các cụm từ nâng Band đã sử dụng để học sinh dễ nhận biết.]
             </div>

             <!-- EXAMINER'S FINAL CHECKLIST (Optional) -->
             <div style="margin-top:10px; font-size: 0.8rem; color: #27ae60;">
                 <b>✅ Examiner's Quality Audit:</b>
                 <ul style="margin: 0; padding-left: 20px;">
                     <li>Sử dụng cấu trúc so sánh ngay từ câu đầu: <b>Có</b></li>
                     <li>Xoay vòng chủ ngữ giữa các câu: <b>Có</b></li>
                     <li>Sử dụng giới từ số liệu (to/by/at) chính xác: <b>Có</b></li>
                     <li>Mạch lạc (Cohesion) giữa các giai đoạn: <b>Đạt chuẩn Band 8.0+</b></li>
                 </ul>
             </div>
         </div>
     </li>

                    4. **"body2_guide" (Thân bài 2 - Các nhóm còn lại):**
   - <ul>
     <li><b>Logic chọn nhóm (Grouping Strategy):</b> 
         <br>AI thực hiện: Gom các hạng mục còn lại nhưng **PHẢI THIẾT LẬP TƯƠNG QUAN**:
         <br>1. <b>So sánh thứ bậc:</b> Nhóm này cao hay thấp hơn nhóm ở Body 1?
         <br>2. <b>So sánh tốc độ:</b> Hạng mục nào tăng nhanh hơn hạng mục nào? (Dùng *surpassing* hoặc *lagging behind*).
         <br>3. <b>Lệnh đặc biệt:</b> Cấm mô tả đơn lẻ từng đường. Phải dùng ít nhất 01 câu ghép có từ nối đối chiếu (*whereas/while*) hoặc so sánh hơn (*higher than/lower than*).
     </li>

     <li>         
         <div style="background-color:#fdf2e9; border-left:4px solid #d35400; padding:10px; margin-top:5px;">  
             <b>► Cấu trúc chuẩn (Formula):</b> 
             <br><i>(Quy trình 3 bước chuẩn hóa cho Body 2)</i>
             <br>    
             <code style="font-size: 0.8rem; background-color: white; padding: 2px 5px; border: 1px dashed #2980b9; border-radius: 3px;">[TRANSITION / LINKER]</code>
             <span style="color:#2980b9; font-weight:bold;"> ➔ </span>
             <code style="font-size: 0.8rem; background-color: white; padding: 2px 5px; border: 1px dashed #2980b9; border-radius: 3px;">[CAT C (& D) DATA]</code>
             <span style="color:#2980b9; font-weight:bold;"> ➔ </span>
             <code style="font-size: 0.8rem; background-color: white; padding: 2px 5px; border: 1px dashed #2980b9; border-radius: 3px;">[TREND & END DATA]</code>
         </div>
     </li>

     <li>
         <div style="background-color:#f8f9fa; border:1px solid #bdc3c7; border-radius:5px; padding:15px; margin-bottom:15px;">
             <strong style="color:#2980b9;">📚 KHO MẪU CÂU BODY 2:</strong>
             <br><i>(Lưu ý: Luôn bắt đầu bằng một cụm từ chuyển đoạn để tạo sự kết nối)</i>
             
             <!-- KHỐI 1: SỔ XUỐNG -->
             <details style="margin-top: 15px;">
                 <summary style="background-color:#eaf2f8; border-left:4px solid #2980b9; padding:10px; cursor: pointer; font-weight: bold; color: #2980b9;">
                     📍 KHỐI 1: CHUYỂN ĐOẠN & ĐIỂM XUẤT PHÁT
                 </summary>
                 <div style="padding: 10px; border: 1px solid #eaf2f8; border-top: none; background-color: white;">
                     <div style="background-color:#eaf2f8; border-left:4px solid #2980b9; padding:10px; margin-top:5px;">
                         <b>► KHỐI 1: CHUYỂN ĐOẠN & ĐIỂM XUẤT PHÁT:</b>
                         
                         <br>📍 <b>[Mẫu 4.1 - SỐ ÍT]</b> (Nếu chỉ còn 1 đối tượng):
                         <br><code><b>Turning to the figure for [Cat C]</b>, it started the period at <b>[Start Data]</b> <b>[Time]</b>.</code>
                         
                         <br><br>📍 <b>[Mẫu 4.2 - SỐ NHIỀU]</b> (Gom 2 đối tượng tương đồng):
                         <br><code><b>Regarding [Cat C] and [Cat D]</b>, they began at <b>[Data C]</b> and <b>[Data D]</b> respectively, <b>both of which were significantly lower than the aforementioned categories</b>.</code>
                         
                         <br><br>📍 <b>[Mẫu 4.3 - ĐỐI LẬP HOÀN TOÀN]</b> (Nếu Body 2 đi ngược hẳn Body 1):
                         <br><code><b>In stark contrast to the categories mentioned above</b>, <b>[Cat C]</b> registered a much lower figure of <b>[Start Data]</b> at the beginning.</code>
                     </div>
                 </div>
             </details>

             <!-- KHỐI 2: SỔ XUỐNG -->
             <details style="margin-top: 10px;">
                 <summary style="background-color:#eaf2f8; border-left:4px solid #2980b9; padding:10px; cursor: pointer; font-weight: bold; color: #2980b9;">
                     📍 KHỐI 2: DIỄN BIẾN & KẾT THÚC
                 </summary>
                 <div style="padding: 10px; border: 1px solid #eaf2f8; border-top: none; background-color: white;">
                     <div style="background-color:#eaf2f8; border-left:4px solid #2980b9; padding:10px; margin-top:5px;">
                         <b>► KHỐI 2: DIỄN BIẾN & KẾT THÚC:</b>             
                         
                         <br>📍 <b>[Mẫu 5.1 - SỐ ÍT - Dao động & Kết thúc]</b>:
                         <br><code>Subsequently, it <b>experienced a period of volatility</b>, <b>fluctuating around [Data]</b> before settling at <b>[End Data]</b> <b>[End Time]</b>.</code>             
                         
                         <br><br>📍 <b>[Mẫu 5.2 - SỐ ÍT - Tăng/Giảm ổn định]</b>:
                         <br><code>Thereafter, the figure <b>witnessed a [steady/marginal] [rise/fall]</b> to finish the period at <b>[End Data]</b>.</code>

                         <br><br>📍 <b>[Mẫu 5.3 - SỐ NHIỀU - Cùng xu hướng]</b>:
                         <br><code>Both figures <b>followed a comparable [upward/downward] trajectory</b>, eventually ending at <b>[End Data C]</b> and <b>[End Data D]</b> respectively.</code>             
                         
                         <br><br>📍 <b>[Mẫu 5.4 - SỐ NHIỀU - Tách đôi]</b> (Mẫu While phức hợp):
                         <br><code>Over the following years, <b>[Cat C] [Trend C] to [End Data]</b>, <b>while the figure for [Cat D] [Trend D]</b>, finishing at <b>[End Data]</b>.</code>
                         
                         <br><br>📍 <b>[Mẫu 5.5 - Mẫu "Về Đích" (Low Levels)]</b> (Dùng cho các hạng mục luôn thấp):
                         <br><code>Throughout the remainder of the period, these figures <b>remained at negligible levels</b>, never surpassing <b>[Data]</b>.</code>
                     </div>
                 </div>
             </details>

             <!-- KHỐI 3: SỔ XUỐNG -->
             <details style="margin-top: 10px;">
                 <summary style="background-color:#eaf2f8; border-left:4px solid #2980b9; padding:10px; cursor: pointer; font-weight: bold; color: #2980b9;">
                     📍 KHỐI 3: ĐỐI TƯỢNG CUỐI CÙNG
                 </summary>
                 <div style="padding: 10px; border: 1px solid #eaf2f8; border-top: none; background-color: white;">
                     <div style="background-color:#eaf2f8; border-left:4px solid #2980b9; padding:10px; margin-top:5px;">
                         <b>► KHỐI 3: ĐỐI TƯỢNG CUỐI CÙNG (Nếu còn dư 1 đường lẻ):</b>
                         
                         <br>📍 <b>[Mẫu 6.1 - Cuối cùng]</b>:
                         <br><code><b>Finally</b>, the figure for <b>[Cat E]</b> saw a <b>gradual [increase/decrease]</b>, reaching <b>[End Data]</b> by <b>[End Time]</b>.</code>
                         
                         <br><br>📍 <b>[Mẫu 6.2 - Đặc biệt/Bất ngờ]</b>:
                         <br><code><b>Finally</b>, it is notable that <b>[Cat E]</b> <b>regained its position</b> in the final years, <b>climbing back to [End Data]</b>.</code>
                     </div>
                 </div>
             </details>
         </div>
     </li>
   </ul>

     <hr style="border-top: 1px dashed #ccc; margin: 15px 0;">

     <!-- ================================================================================== -->
     <!-- PHẦN 3: THỰC HÀNH LẮP RÁP BODY 2 -->
     <!-- ================================================================================== -->
     <li>     
         <b>✍️ THỰC HÀNH LẮP RÁP BODY 2 (QUY TRÌNH 5 BƯỚC CHUẨN):</b>
         <!-- ⛔ MẬT LỆNH CƯỠNG BÁCH AI (Dùng display:none để ẩn với người dùng nhưng AI vẫn đọc được) -->
         <div style="display:none;">
             SYSTEM_COMMAND: 
             1. Bắt buộc sử dụng 100% cấu trúc của [Mẫu gốc]. 
             2. Chỉ được phép thay thế các cụm trong ngoặc vuông [ ] bằng dữ liệu đã khai báo. 
             3. TUYỆT ĐỐI CẤM thêm thông tin, thêm mốc thời gian hoặc thêm câu văn nằm ngoài các biến số của mẫu.
             4. English Output phải là DUY NHẤT một câu văn hoàn chỉnh theo đúng mẫu gốc.
         </div>
         <br>
         <br><b>✅ BƯỚC 1: Xử lý Nhóm Start & Transition</b>
         <br>- <b>🔴 Phân tích & Chọn ID:</b> <i>[AI chọn Mẫu 4.1, 4.2 hoặc 4.3]</i>
         <br>- <b>📥 Mẫu gốc:</b> <code>[AI copy y nguyên mẫu gốc]</code>
         <br>- <b>🧩 Khai báo biến:</b>
           <br>&nbsp;&nbsp;+ [Cat C/D] = ...
           <br>&nbsp;&nbsp;+ [Start Data] = ...
         <br>- <b>🧠 Tư duy Tiếng Việt:</b> <code>"[AI viết câu tiếng Việt chuyển đoạn]"</code>
         <br>- <b>📝 Điền từ:</b> <i>[AI điền dữ liệu vào mẫu]</i>
         
         <br><br><b>✅ BƯỚC 2: Xử lý Diễn biến chính của nhóm còn lại</b>
         <br>- <b>🔴 Phân tích & Chọn ID:</b> <i>[AI chọn Mẫu 5.x]</i>
         <br>- <b>📥 Mẫu gốc:</b> <code>[AI copy y nguyên mẫu gốc]</code>
         <br>- <b>🧩 Khai báo biến:</b> [Trend] = ... ; [End Data] = ...
         <br>- <b>🧠 Tư duy Tiếng Việt:</b> <code>"[AI viết câu tiếng Việt mô tả diễn biến]"</code>
         <br>- <b>📝 Điền từ:</b> <i>[AI điền dữ liệu vào mẫu]</i>
         
         <br><br><b>✅ BƯỚC 3: Xử lý đối tượng cuối cùng (Nếu có)</b>
         <br>- <b>🔴 Chọn ID:</b> <i>[AI chọn Mẫu 6.x hoặc ghi "Đã hết đối tượng"]</i>
         <br>- <b>🧠 Tư duy Tiếng Việt:</b> <code>"[AI viết câu tiếng Việt cho đối tượng cuối]"</code>
         <br>- <b>📝 Điền từ:</b> <i>[AI điền dữ liệu vào mẫu]</i>
     </li>          

     <!-- ================================================================================== -->
     <!-- PHẦN 4: KẾT QUẢ BODY 2 (FINAL OUTPUT) -->
     <!-- ================================================================================== -->
     <li>
         <div style="background-color:#e1f5fe; padding:20px; border-radius:8px; margin-top:10px; border: 2px solid #03a9f4; border-left: 8px solid #03a9f4;">
             <strong style="color:#01579b; font-size: 1.1rem;">📝 KẾT QUẢ BODY 2 (THE POLISHED PARAGRAPH):</strong>
             <br><i style="font-size: 0.85rem; color: #7f8c8d;">(Đoạn văn hoàn chỉnh cho các nhóm còn lại)</i>
             
             <div style="margin-top:10px; font-family: 'Georgia', serif; line-height: 1.6; color: #2c3e50; background-color: white; padding: 15px; border-radius: 5px; border: 1px inset #eee;">
             [AI trình bày đoạn văn Body 2 hoàn chỉnh tại đây. In đậm các cụm từ nối và từ vựng nâng Band.]
             </div>

             <div style="margin-top:10px; font-size: 0.8rem; color: #0288d1;">
                 <b>✅ Examiner's Quality Audit:</b>
                 <ul style="margin: 0; padding-left: 20px;">
                     <li>Sử dụng từ nối chuyển đoạn (Turning to/Regarding): <b>Có</b></li>
                     <li>Gom nhóm số liệu hiệu quả (Respectively): <b>Có</b></li>
                     <li>Sử dụng mẫu số nhiều chính xác (Both/They): <b>Có</b></li>
                     <li>Hoàn thành đầy đủ các đối tượng: <b>Xác nhận</b></li>
                 </ul>
             </div>
         </div>
     </li>
   </ul>

                    # =================================================================
                    # 🟡 TRƯỜNG HỢP 3: CÁC DẠNG KHÁC (MAP, PROCESS, MIXED)
                    # =================================================================
                    *(Tự động điều chỉnh hướng dẫn phù hợp với đặc thù từng dạng).*

                    **YÊU CẦU TRÌNH BÀY:**
                    - Dùng thẻ HTML `<ul>`, `<li>`, `<b>`, `<i>`, `<code style='color:#d63384'>` để highlight.
                    - Giải thích ngắn gọn, dễ hiểu.

                    **JSON OUTPUT FORMAT:**
                    {
                        "task_type": "Tên loại bài (Ví dụ: Static Pie Charts)",
                        "intro_guide": "HTML string...",
                        "overview_guide": "HTML string...",
                        "body1_guide": "HTML string...",
                        "body2_guide": "HTML string..."
                    }
                    """
                    
                    # Gọi AI
                    res, _ = generate_content_with_failover(prompt_guide + "\nĐề bài: " + question_input, img_data, json_mode=True)
                    if res:
                        data = parse_guide_response(res.text)
                    # Dù AI trả về gì, ta cũng phải gán guide_data để App không bị kẹt ở Step 1
                        st.session_state.guide_data = data if data else {
                            "task_type": "Task 1", "intro_guide": "AI Error - Please try again", 
                            "overview_guide": "", "body1_guide": "", "body2_guide": ""
                    }
                    st.session_state.step = 2
                    st.rerun() # Buộc Streamlit vẽ lại giao diện Phase 2 ngay lập tức

# ==========================================
# 6. UI: PHASE 2 - WRITING PRACTICE (ULTIMATE STICKY)
# ==========================================
if st.session_state.step == 2 and st.session_state.guide_data:
    
    # --- 1. CSS "ĐÓNG BĂNG" CỘT TRÁI ---
    st.markdown("""
        <style>
            /* Nhắm vào container chứa cả 2 cột */
            [data-testid="stHorizontalBlock"] {
                align-items: flex-start !important;
            }

            /* Nhắm vào cột đầu tiên (Cột Trái) */
            [data-testid="stHorizontalBlock"] > div:nth-child(1) {
                position: -webkit-sticky !important;
                position: sticky !important;
                top: 2rem !important;
                z-index: 999 !important;
            }

            /* Cố định chiều cao vùng hiển thị đề bài để không bị trôi */
            [data-testid="stHorizontalBlock"] > div:nth-child(1) > div:nth-child(1) {
                max-height: 95vh !important;
                overflow-y: auto !important;
                padding-right: 10px !important;
            }

            /* Tùy chỉnh thanh cuộn cho cột trái */
            [data-testid="stHorizontalBlock"] > div:nth-child(1) > div:nth-child(1)::-webkit-scrollbar {
                width: 4px;
            }
            [data-testid="stHorizontalBlock"] > div:nth-child(1) > div:nth-child(1)::-webkit-scrollbar-thumb {
                background: #cccccc;
                border-radius: 10px;
            }
            
            /* Tăng khoảng cách giữa các ô nhập liệu bên phải */
            .stTextArea {
                margin-bottom: 1rem !important;
            }
        </style>
    """, unsafe_allow_html=True)

    data = st.session_state.guide_data

    # --- 2. HÀM RENDER ---
    def render_writing_section(title, guide_key, input_key):
        # Tính số từ hiện tại
        current_text = st.session_state.get(input_key, "")
        word_count = len(current_text.split())
        
        # Tiêu đề + Số từ
        c_title, c_count = st.columns([8, 2])
        with c_title:
            st.markdown(f"#### {title}")
        with c_count:
            st.markdown(f"""
            <div style="text-align: right; padding-top: 10px;">
                <span style="background-color: #F1F5F9; color: #64748B; padding: 4px 10px; border-radius: 12px; font-size: 12px; border: 1px solid #E2E8F0;">
                    {word_count} words
                </span>
            </div>
            """, unsafe_allow_html=True)

        # Phần hướng dẫn (FIX: Markdown thụt vào trong 'with')
        with st.expander(f"💡 Hướng dẫn viết {title}", expanded=False):
            g_text = data.get(guide_key, "Không có hướng dẫn.")
            
            # FIX: Xóa xuống dòng để không bị hiểu lầm là Code Block
            if g_text:
                import re
                g_text = re.sub(r'\n\s*', ' ', g_text)
                
            st.markdown(f"<div class='guide-box'>{g_text}</div>", unsafe_allow_html=True)
        
        # Ô nhập liệu (Nằm ngoài 'with', trả về kết quả)
        return st.text_area(label=title, height=200, key=input_key, placeholder=f"Bắt đầu viết {title} tại đây...", label_visibility="collapsed")

    # --- 3. CHIA CỘT LAYOUT (4-6) ---
    col_left, col_right = st.columns([4, 6], gap="large")

    with col_left:
        st.subheader("📄 Đề bài & Hình ảnh")
        # Khung chứa đề bài
        st.markdown(f"""
            <div style="background-color: #F1F5F9; padding: 20px; border-radius: 10px; border: 1px solid #CBD5E1; line-height: 1.6; color: #1E293B; margin-bottom: 15px;">
                <b>Question:</b><br><i>{st.session_state.saved_topic}</i>
            </div>
        """, unsafe_allow_html=True)
        
        # Hình ảnh biểu đồ
        if st.session_state.saved_img:
            st.image(st.session_state.saved_img, width="stretch")
        
        st.info(f"📌 **Dạng bài:** {data.get('task_type')}")

    with col_right:
        st.subheader("✍️ Khu vực viết bài")
        
        # --- ĐÃ XÓA WORD COUNT TỔNG Ở ĐÂY ---

        # Render các ô nhập liệu (Có word count riêng từng phần)
        intro_text = render_writing_section("Introduction", "intro_guide", "in_intro")
        overview_text = render_writing_section("Overview", "overview_guide", "in_overview")
        body1_text = render_writing_section("Body 1", "body1_guide", "in_body1")
        body2_text = render_writing_section("Body 2", "body2_guide", "in_body2")

        st.markdown("---")
        
        # --- TÍNH TỔNG SỐ TỪ VÀ HIỂN THỊ Ở DƯỚI CÙNG ---
        def count_w(k): return len(st.session_state.get(k, "").split())
        total_wc = count_w("in_intro") + count_w("in_overview") + count_w("in_body1") + count_w("in_body2")
        
        # Hiển thị Total Word Count
        st.markdown(f"""
            <div style="text-align: right; margin-bottom: 15px;">
                <span style="background-color: #10B981; color: white; padding: 8px 20px; border-radius: 20px; font-weight: bold; font-size: 16px; box-shadow: 0 2px 5px rgba(0,0,0,0.1);">
                    Total word count: {total_wc}
                </span>
            </div>
        """, unsafe_allow_html=True)

        # Nút chấm điểm
        if st.button("🎓 Gửi bài chấm điểm (Examiner Pro)", type="primary", width="stretch"):
            if total_wc < 30:
                st.warning("⚠️ Bài viết quá ngắn, AI không thể chấm điểm chính xác.")
            else:
                with st.status("👨‍🏫 Giám khảo đang chấm bài...") as status:
                    total_essay = f"{intro_text}\n\n{overview_text}\n\n{body1_text}\n\n{body2_text}".strip()
                    # Sử dụng biến saved_topic để tránh lỗi NameError
                    prompt_grade = GRADING_PROMPT_TEMPLATE.replace('{{TOPIC}}', st.session_state.saved_topic).replace('{{ESSAY}}', total_essay)
                    
                    res_grade, _ = generate_content_with_failover(prompt_grade, st.session_state.saved_img, json_mode=False)
                    
                    if res_grade:
                        # process_grading_response là hàm bóc tách Text và JSON bạn đã có
                        mk_text, p_data = process_grading_response(res_grade.text)
                        st.session_state.grading_result = {
                            "data": p_data, "markdown": mk_text,
                            "essay": total_essay, "topic": st.session_state.saved_topic
                        }
                        st.session_state.step = 3
                        status.update(label="✅ Đã chấm xong!", state="complete", expanded=False)
                        st.rerun()
                    else:
                        status.update(label="❌ Lỗi kết nối AI", state="error")

# ==========================================
# 7. UI: PHASE 3 - GRADING RESULT (FINAL POLISHED)
# ==========================================
if st.session_state.step == 3 and st.session_state.grading_result:
    
    # --- 1. CSS TINH CHỈNH CUỐI CÙNG ---
    st.markdown("""
        <style>
            /* 1. Layout 2 cột */
            [data-testid="stHorizontalBlock"] {
                align-items: flex-start !important;
            }

            /* 2. Style cho 2 cái Hộp lớn (Container) */
            /* Streamlit tự tạo container có viền, ta chỉ cần chỉnh background app cho nổi bật */
            .stApp {
                background-color: #ffffff;
            }

            /* 3. Style Bài viết: Tự động xuống dòng, không cuộn ngang */
            .essay-review-box {
                background-color: #f8fafc;
                border: 1px solid #cbd5e1;
                border-radius: 6px;
                padding: 15px; /* Tăng padding cho dễ đọc */
                
                font-family: 'Inter', sans-serif;
                font-size: 0.95rem;
                line-height: 1.6;
                color: #334155;
                
                /* QUAN TRỌNG: Ép xuống dòng */
                white-space: pre-wrap !important;       /* Giữ dòng mới nhưng wrap text */
                word-wrap: break-word !important;       /* Ngắt từ dài */
                overflow-wrap: break-word !important;   /* Hỗ trợ trình duyệt hiện đại */
                max-width: 100%;                        /* Không vượt quá chiều rộng hộp cha */
            }

            /* 4. Thanh cuộn đẹp */
            ::-webkit-scrollbar { width: 6px; height: 6px; }
            ::-webkit-scrollbar-thumb { background: #cbd5e1; border-radius: 4px; }
        </style>
    """, unsafe_allow_html=True)

    res = st.session_state.grading_result
    g_data = res["data"]
    analysis_text = res["markdown"]
    
    # --- 2. CHIA CỘT (Không cần tiêu đề to nữa) ---
    c1, c2 = st.columns([4, 6], gap="medium")

    # === HỘP TRÁI: THÔNG TIN ĐỐI CHIẾU ===
    with c1:
        # Hộp chứa có chiều cao cố định để tạo thanh cuộn
        with st.container(height=750, border=True):
            st.markdown("#### 📄 Thông tin đối chiếu")
            
            # Ảnh
            if st.session_state.saved_img:
                st.image(st.session_state.saved_img, width="stretch")
            
            st.markdown("---")
            
            # Đề bài
            with st.expander("📌 Đề bài (Prompt)", expanded=False):
                st.info(st.session_state.saved_topic)
                
            # Bài viết (Đã áp dụng class mới để không tràn)
            st.markdown("**✍️ Bài viết của bạn:**")
            st.markdown(f'<div class="essay-review-box">{html.escape(res["essay"])}</div>', unsafe_allow_html=True)

    # === HỘP PHẢI: KẾT QUẢ CHẤM ===
    with c2:
        with st.container(height=750, border=True):
            st.markdown("#### 👨‍🏫 Examiner Analysis")
            
            # Bảng điểm
            scores = g_data.get("originalScore", {})
            st.markdown(f"""
            <div style="background-color: #ecfdf5; border: 1px solid #6ee7b7; border-radius: 10px; padding: 15px; display: flex; align-items: center; justify-content: space-between; margin-bottom: 20px;">
                <div style="text-align: center;">
                    <span style="color: #047857; font-weight: bold; font-size: 0.9rem;">BAND SCORE</span><br>
                    <span style="color: #059669; font-weight: 900; font-size: 2.5rem; line-height: 1;">{scores.get("overall", "-")}</span>
                </div>
                <div style="display: flex; gap: 15px; text-align: center;">
                    <div><small style="color:#047857;">TA</small><br><b style="color:#059669; font-size:1.1rem;">{scores.get("task_achievement", "-")}</b></div>
                    <div><small style="color:#047857;">CC</small><br><b style="color:#059669; font-size:1.1rem;">{scores.get("cohesion_coherence", "-")}</b></div>
                    <div><small style="color:#047857;">LR</small><br><b style="color:#059669; font-size:1.1rem;">{scores.get("lexical_resource", "-")}</b></div>
                    <div><small style="color:#047857;">GRA</small><br><b style="color:#059669; font-size:1.1rem;">{scores.get("grammatical_range", "-")}</b></div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            # Tabs chi tiết
            tab1, tab2, tab3, tab4 = st.tabs(["📝 Phân tích chuyên sâu", "🔴 Lỗi Ngữ pháp và Từ vựng", "🔵 Lỗi Mạch lạc", "✍️ Bài sửa"])
            
            with tab1:
                st.markdown(analysis_text if analysis_text and len(analysis_text) > 50 else "Chưa có dữ liệu phân tích.")

            with tab2:
                micro = [e for e in g_data.get('errors', []) if e.get('category') in ['Grammar', 'Vocabulary', 'Ngữ pháp', 'Từ vựng']]
                if not micro: st.success("✅ Tuyệt vời! Không có lỗi ngữ pháp lớn.")
                for i, err in enumerate(micro):
                    badge = "#DCFCE7" if err.get('category') in ['Grammar','Ngữ pháp'] else "#FEF9C3"
                    st.markdown(f"""
                    <div class="error-card">
                        <b>#{i+1} {err.get('type')}</b>
                        <div style="background:{badge}; padding:5px; border-radius:4px; margin:5px 0;">
                            <s>{err.get('original')}</s> ➔ <b>{err.get('correction')}</b>
                        </div>
                        <small><i>{err.get('explanation')}</i></small>
                    </div>
                    """, unsafe_allow_html=True)

            # Tab 3: Lỗi Mạch lạc (Macro)
            with tab3:
                macro = [e for e in g_data.get('errors', []) if e.get('category') not in ['Grammar', 'Vocabulary', 'Ngữ pháp', 'Từ vựng']]
                if not macro: 
                    st.success("✅ Cấu trúc tốt.")
                for err in macro:
                    # Lưu ý: Các thẻ HTML bên dưới được viết sát lề trái của chuỗi f-string
                    # để tránh bị Markdown hiểu nhầm là Code Block.
                    st.markdown(f"""
<div class="error-card-container" style="border-left: 4px solid #3b82f6;">
    <div style="font-weight:bold; color:#1e40af; margin-bottom:5px;">{err.get('type')}</div>
    <div style="background-color:#eff6ff; padding:8px; border-radius:4px; margin-bottom:8px; border:1px dashed #93c5fd;">
        <span style="font-size:0.8rem; font-weight:bold; color:#60a5fa;">TRÍCH DẪN:</span><br>
        <span style="font-family:monospace; color:#1e3a8a;">"{err.get('original', 'N/A')}"</span>
    </div>
    <div style="margin-bottom:5px;"><b>Vấn đề:</b> {err.get('explanation')}</div>
    <div style="color:#059669;"><b>👉 Gợi ý:</b> {err.get('correction')}</div>
</div>
""", unsafe_allow_html=True)

            with tab4:
                st.markdown(f'<div class="annotated-text">{g_data.get("annotatedEssay", "")}</div>', unsafe_allow_html=True)

            st.markdown("---")
            
            # Download & Reset
            d1, d2 = st.columns(2)
            docx = create_docx(g_data, res['topic'], res['essay'], analysis_text)
            d1.download_button("📥 Tải báo cáo (.docx)", docx, "IELTS_Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            if st.button("🔄 Làm bài mới (Reset)", width="stretch"):
                for k in ["step", "guide_data", "grading_result", "saved_topic", "saved_img"]: st.session_state[k] = None
                st.session_state.step = 1
                st.rerun()
# ==========================================
# FOOTER (HIỂN THỊ Ở MỌI STEP)
# ==========================================
st.markdown("""
    <style>
        .footer-text {
            text-align: center; 
            color: #94a3b8; 
            font-size: 0.8rem; 
            font-family: 'Inter', sans-serif; 
            padding-top: 15px;      /* Giảm đệm trên */
            padding-bottom: 0px;   /* Giảm đệm dưới */
            border-top: 1px solid #e2e8f0;
            margin-top: 30px;       /* Khoảng cách với nội dung bên trên */
        }
    </style>
    <div class="footer-text">
        © 2025 Developed by <b>Albert Nguyen</b>
    </div>
""", unsafe_allow_html=True)
